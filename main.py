# -*- coding: utf-8 -*-
"""
App Streamlit para generar textos de etiquetas y fichas técnicas de piensos.

Ejecutar:
    streamlit run main.py

Objetivo operativo:
- Cargar un fichero de formulación en formato Excel, CSV, TXT, JSON, PDF o DOCX.
- Detectar productos y datos de fórmula/nutrientes.
- Cargar textos parametrizados desde un Excel con hojas "Etiquetas" y "Beneficios".
- Seleccionar producto, especie, subespecie, lifestage y textos propuestos.
- Editar todos los textos antes de exportar.
- Generar: Etiqueta, FT Calidad-Operaciones, FT Comercial y FT Especificaciones.
- Descargar cada ficha en TXT, DOCX, PDF y XLSX.
- Acumular resultados y exportarlos masivamente a Excel.

Nota técnica:
No existe un lector realmente universal para "cualquier formato". Esta app implementa
lectura robusta para formatos habituales y un parser específico para salidas de formulación
tipo Multi-Mix en texto de una columna, como el fichero de ejemplo aportado.
"""

from __future__ import annotations

import io
import json
import re
import base64
import html as html_lib
import unicodedata
from datetime import date
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

try:
    from charset_normalizer import from_bytes as charset_from_bytes
except Exception:  # pragma: no cover
    charset_from_bytes = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from xml.sax.saxutils import escape as xml_escape
except Exception:  # pragma: no cover
    A4 = None
    getSampleStyleSheet = None
    cm = None
    colors = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None
    Table = None
    TableStyle = None
    xml_escape = None


APP_TITLE = "Generador de etiquetas y fichas técnicas de piensos"

DOC_TYPES = [
    "Etiqueta",
    "FT Calidad-Operaciones",
    "FT Comercial",
    "FT Especificaciones",
]

QUALITY_OPERATION_FIELDS = [
    "Fecha",
    "Versión",
    "Nombre comercial",
    "Código Unite",
    "EAN",
    "Tipo de pienso",
    "Especie",
    "Subespecie",
    "Lifestage",
    "Animales de destino",
    "Modo de empleo",
    "Precauciones de uso",
    "Recomendaciones de manejo en etiquetado",
    "Imagen envase",
    "Peso del producto en saco",
    "Ficha técnica del envase",
    "Formato de palé. Mosaico, alturas y peso",
    "Características nutricionales",
    "Materias primas, aditivos y correctores específicos",
    "Analíticas especiales además del plan analítico",
    "Presentación",
    "Durabilidad mínima %",
    "Alerta durabilidad %",
    "Finos máximos %",
    "Alerta finos %",
    "Homologado",
    "Medicado",
    "Periodo de espera",
    "Fábrica",
    "Dirección",
    "Código postal",
    "Población",
    "Provincia",
    "Teléfono",
]

COMMERCIAL_FIELDS = [
    "Nombre comercial",
    "Especie",
    "Subespecie",
    "Lifestage",
    "Animales de destino",
    "Tipo de pienso",
    "Definición / Posicionamiento",
    "Características",
    "Foco-beneficio",
    "Beneficios",
    "Modo de empleo",
    "Precauciones de uso",
    "Presentación",
    "Peso del producto en saco",
]

LABEL_FIELDS = [
    "Nombre comercial",
    "Tipo de pienso",
    "Especie",
    "Subespecie",
    "Animales de destino",
    "Modo de empleo",
    "Precauciones de uso",
    "Recomendaciones de manejo en etiquetado",
    "Presentación",
    "Peso del producto en saco",
    "EAN",
]

SPECIFICATION_FIELDS = COMMERCIAL_FIELDS + [
    "Código Unite",
    "EAN",
    "Homologado",
    "Medicado",
    "Periodo de espera",
    "Fábrica",
    "Características nutricionales",
    "Materias primas, aditivos y correctores específicos",
    "Analíticas especiales además del plan analítico",
    "Durabilidad mínima %",
    "Alerta durabilidad %",
    "Finos máximos %",
    "Alerta finos %",
]

ALL_FIELD_OPTIONS = list(dict.fromkeys(
    LABEL_FIELDS + QUALITY_OPERATION_FIELDS + COMMERCIAL_FIELDS + SPECIFICATION_FIELDS
))

DEFAULT_NUTRIENT_HINTS = [
    "PROT_BRU",
    "PROTEINA",
    "PROTEÍNA",
    "GRASA_BR",
    "FIBRA_BR",
    "CENIZAS",
    "ALM_EWER",
    "ALMIDON",
    "AZUCARES",
    "A+A",
    "CA",
    "P_",
    "NA",
    "CU",
    "LYS",
    "LISINA",
    "MET",
    "THR",
    "TRP",
    "NE_SW",
    "ME_SW",
    "UFL",
    "UFV",
    "PDIE",
    "PDIN",
    "PDIA",
]


BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
ASSETS_DIR = BASE_DIR / "assets"
TEMPLATES_DIR = BASE_DIR / "templates"

MASTER_TEXTS_PATH = DATA_DIR / "Etiquetas_y_beneficios_para_App_animales_destino.xlsx"
NUTRIENT_DEFAULTS_PATH = BASE_DIR / "nutrient_defaults.json"
QUALITY_OPERATION_DEFAULTS_PATH = BASE_DIR / "quality_operation_defaults.json"

NUTRIENT_PROFILES = {
    "commercial": "FT Comercial",
    "technical": "FT Calidad-Operaciones y FT Especificaciones",
}

DOC_NUTRIENT_PROFILE = {
    "FT Comercial": "commercial",
    "FT Calidad-Operaciones": "technical",
    "FT Especificaciones": "technical",
}




# ---------------------------------------------------------------------------
# Utilidades generales
# ---------------------------------------------------------------------------

def strip_accents(value: Any) -> str:
    text = "" if value is None else str(value)
    text = unicodedata.normalize("NFKD", text)
    return "".join(ch for ch in text if not unicodedata.combining(ch))


def norm_key(value: Any) -> str:
    text = strip_accents(value).lower().strip()
    return re.sub(r"[^a-z0-9]+", "", text)


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    text = str(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def unique_clean(values: Iterable[Any]) -> List[str]:
    out: List[str] = []
    seen = set()
    for value in values:
        text = clean_text(value)
        if not text:
            continue
        key = norm_key(text)
        if key not in seen:
            out.append(text)
            seen.add(key)
    return out


def decode_bytes(data: bytes) -> str:
    if not data:
        return ""
    if charset_from_bytes is not None:
        try:
            result = charset_from_bytes(data).best()
            if result:
                return str(result)
        except Exception:
            pass
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("latin-1", errors="replace")


def is_numeric_text(value: Any) -> bool:
    text = clean_text(value).replace(",", ".")
    if text in {"", "."}:
        return False
    try:
        float(text)
        return True
    except Exception:
        return False


def parse_float(value: Any) -> Optional[float]:
    text = clean_text(value).replace(",", ".")
    if not text or text == ".":
        return None
    try:
        return float(text)
    except Exception:
        return None


def df_to_raw_text(df: pd.DataFrame) -> str:
    """Convert a dataframe to raw text preserving spacing inside cells.

    This is important for formulation outputs exported to Excel as a single
    text column: fixed-width spacing is needed to parse ingredient and analysis
    tables reliably.
    """
    lines = []
    for _, row in df.iterrows():
        vals = []
        for value in row.tolist():
            if value is None:
                continue
            if isinstance(value, float) and pd.isna(value):
                continue
            raw = str(value).replace("\r\n", "\n").replace("\r", "\n").rstrip()
            if raw.strip():
                vals.append(raw)
        if vals:
            lines.append(" ".join(vals))
    return "\n".join(lines)


def safe_filename(name: str) -> str:
    base = strip_accents(name)
    base = re.sub(r"[^A-Za-z0-9._-]+", "_", base).strip("_")
    return base[:120] or "documento"


def load_json_safe(path: Path, default: Any) -> Any:
    try:
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default
    return default


def save_json_safe(path: Path, data: Any) -> Tuple[bool, str]:
    try:
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return True, f"Información guardada en {path.name}."
    except Exception as exc:
        return False, f"No se ha podido guardar en {path.name}: {exc}"


def load_nutrient_defaults() -> Dict[str, List[str]]:
    raw = load_json_safe(NUTRIENT_DEFAULTS_PATH, {})
    if not isinstance(raw, dict):
        raw = {}
    return {profile: unique_clean(raw.get(profile, [])) for profile in NUTRIENT_PROFILES}


def save_nutrient_defaults(defaults: Dict[str, List[str]]) -> Tuple[bool, str]:
    safe_defaults = {profile: unique_clean(defaults.get(profile, [])) for profile in NUTRIENT_PROFILES}
    return save_json_safe(NUTRIENT_DEFAULTS_PATH, safe_defaults)


def get_saved_nutrient_defaults() -> Dict[str, List[str]]:
    if "saved_nutrient_defaults" not in st.session_state:
        st.session_state["saved_nutrient_defaults"] = load_nutrient_defaults()
    return st.session_state["saved_nutrient_defaults"]


def set_saved_nutrient_defaults(profile: str, nutrients: List[str]) -> Tuple[bool, str]:
    defaults = get_saved_nutrient_defaults()
    defaults[profile] = unique_clean(nutrients)
    st.session_state["saved_nutrient_defaults"] = defaults
    return save_nutrient_defaults(defaults)


def filter_available(values: Iterable[str], available: List[str]) -> List[str]:
    available_by_key = {norm_key(value): value for value in available}
    selected: List[str] = []
    for value in values:
        match = available_by_key.get(norm_key(value))
        if match and match not in selected:
            selected.append(match)
    return selected


def default_nutrients_for_profile(profile: str, available: List[str], current_key: str) -> List[str]:
    if current_key in st.session_state:
        return filter_available(st.session_state.get(current_key, []), available)
    saved = filter_available(get_saved_nutrient_defaults().get(profile, []), available)
    if saved:
        return saved
    return default_nutrients(available)


def load_quality_operation_defaults() -> Dict[str, str]:
    raw = load_json_safe(QUALITY_OPERATION_DEFAULTS_PATH, {})
    if not isinstance(raw, dict):
        return {}
    return {clean_text(k): clean_text(v) for k, v in raw.items()}


def save_quality_operation_defaults(values: Dict[str, Any]) -> Tuple[bool, str]:
    safe = {clean_text(k): clean_text(v) for k, v in values.items() if clean_text(k)}
    st.session_state["quality_operation_defaults"] = safe
    return save_json_safe(QUALITY_OPERATION_DEFAULTS_PATH, safe)


def get_quality_operation_defaults() -> Dict[str, str]:
    if "quality_operation_defaults" not in st.session_state:
        st.session_state["quality_operation_defaults"] = load_quality_operation_defaults()
    return st.session_state["quality_operation_defaults"]


QUALITY_WIDGET_FIELDS = {
    "manual_fecha": "Fecha",
    "manual_version": "Versión",
    "manual_codigo_unite": "Código Unite",
    "manual_ean": "EAN",
    "manual_presentacion": "Presentación",
    "manual_peso_saco": "Peso del producto en saco",
    "manual_homologado": "Homologado",
    "manual_medicado": "Medicado",
    "manual_periodo_espera": "Periodo de espera",
    "manual_fabrica": "Fábrica",
    "manual_direccion": "Dirección",
    "manual_codigo_postal": "Código postal",
    "manual_poblacion": "Población",
    "manual_provincia": "Provincia",
    "manual_telefono": "Teléfono",
    "manual_imagen_envase": "Imagen envase",
    "manual_ficha_envase": "Ficha técnica del envase",
    "manual_formato_pale": "Formato de palé. Mosaico, alturas y peso",
    "manual_dur_min": "Durabilidad mínima %",
    "manual_dur_alert": "Alerta durabilidad %",
    "manual_finos_max": "Finos máximos %",
    "manual_finos_alert": "Alerta finos %",
    "manual_definicion": "Definición / Posicionamiento",
    "manual_caracteristicas": "Características",
    "manual_caracteristicas_nut": "Características nutricionales",
    "manual_materias": "Materias primas, aditivos y correctores específicos",
    "manual_analiticas": "Analíticas especiales además del plan analítico",
}


def row_data_lookup(row_data: Dict[str, Any], *names: str) -> str:
    norm_map = {norm_key(k): v for k, v in (row_data or {}).items()}
    for name in names:
        if name in row_data:
            return clean_text(row_data.get(name, ""))
        key = norm_key(name)
        if key in norm_map:
            return clean_text(norm_map.get(key, ""))
    return ""


def init_quality_widget_defaults(row_data: Dict[str, Any], product_data: Dict[str, Any]) -> None:
    saved = get_quality_operation_defaults()
    fallbacks = {
        "Fecha": date.today().isoformat(),
        "Versión": row_data_lookup(row_data, "Versión", "Version") or "1",
        "Código Unite": row_data_lookup(row_data, "Código Unite", "Codigo Unite") or clean_text(product_data.get("product_id", "")),
        "EAN": row_data_lookup(row_data, "EAN"),
        "Presentación": row_data_lookup(row_data, "Presentación", "PRESENTACION"),
        "Peso del producto en saco": row_data_lookup(row_data, "Peso del producto en saco"),
        "Homologado": row_data_lookup(row_data, "Homologado", "HOMOLOGADO"),
        "Medicado": row_data_lookup(row_data, "Medicado", "MEDICADO"),
        "Periodo de espera": row_data_lookup(row_data, "Periodo de espera"),
        "Fábrica": row_data_lookup(row_data, "Fábrica", "FABRICA"),
        "Dirección": row_data_lookup(row_data, "Dirección", "DIRECCION"),
        "Código postal": row_data_lookup(row_data, "Código postal", "CODIGO POSTAL"),
        "Población": row_data_lookup(row_data, "Población", "POBLACION"),
        "Provincia": row_data_lookup(row_data, "Provincia", "PROVINCIA"),
        "Teléfono": row_data_lookup(row_data, "Teléfono", "TELEFONO"),
        "Imagen envase": row_data_lookup(row_data, "Imagen envase"),
        "Ficha técnica del envase": row_data_lookup(row_data, "Ficha técnica del envase"),
        "Formato de palé. Mosaico, alturas y peso": row_data_lookup(row_data, "Formato de palé. Mosaico y alturas y peso", "Formato de palé. Mosaico, alturas y peso"),
        "Durabilidad mínima %": row_data_lookup(row_data, "Durabilidad minima %", "Durabilidad mínima %"),
        "Alerta durabilidad %": row_data_lookup(row_data, "Alerta durabilidad %"),
        "Finos máximos %": row_data_lookup(row_data, "Finos máximos %"),
        "Alerta finos %": row_data_lookup(row_data, "Alerta finos %"),
        "Definición / Posicionamiento": row_data_lookup(row_data, "Definición / Posicionamiento"),
        "Características": row_data_lookup(row_data, "Características"),
        "Características nutricionales": row_data_lookup(row_data, "Características nutricionales"),
        "Materias primas, aditivos y correctores específicos": row_data_lookup(row_data, "Materias primas, aditivos, correctores ESPECIFICOS", "Materias primas, aditivos y correctores específicos"),
        "Analíticas especiales además del plan analítico": row_data_lookup(row_data, "Analíticas especiales además de las ya definidas por plan analítico.", "Analíticas especiales además del plan analítico"),
    }
    for key, field in QUALITY_WIDGET_FIELDS.items():
        if key not in st.session_state:
            st.session_state[key] = clean_text(saved.get(field, "")) or clean_text(fallbacks.get(field, ""))


def collect_quality_values_from_session() -> Dict[str, str]:
    return {field: clean_text(st.session_state.get(key, "")) for key, field in QUALITY_WIDGET_FIELDS.items()}


def asset_to_data_uri(path: Path) -> str:
    if not path.exists():
        return ""
    suffix = path.suffix.lower().replace(".", "")
    mime = "jpeg" if suffix in {"jpg", "jpeg"} else suffix
    try:
        encoded = base64.b64encode(path.read_bytes()).decode("ascii")
        return f"data:image/{mime};base64,{encoded}"
    except Exception:
        return ""


def shorten_for_layout(value: Any, max_chars: int = 360) -> str:
    text = clean_text(value).replace("\n", " ")
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= max_chars:
        return text
    cut = text[:max_chars].rsplit(" ", 1)[0].strip()
    return cut + "…"


def flat_rows_to_field_map(flat_rows: List[Dict[str, Any]]) -> Dict[str, str]:
    data: Dict[str, str] = {}
    for row in flat_rows or []:
        field = clean_text(row.get("Campo", ""))
        value = clean_text(row.get("Valor", ""))
        if field and value and field not in data:
            data[field] = value
    return data


def section_to_html(value: str) -> str:
    text = clean_text(value)
    if not text:
        return ""
    lines = text.splitlines()
    html_parts: List[str] = []
    i = 0
    while i < len(lines):
        table_rows, next_i = collect_markdown_table(lines, i)
        if table_rows:
            html_parts.append("<table>")
            for r, row in enumerate(table_rows):
                tag = "th" if r == 0 else "td"
                html_parts.append("<tr>" + "".join(f"<{tag}>{html_lib.escape(cell)}</{tag}>" for cell in row) + "</tr>")
            html_parts.append("</table>")
            i = next_i
            continue
        line = clean_text(lines[i])
        if line:
            html_parts.append(f"<p>{html_lib.escape(line)}</p>")
        i += 1
    return "\n".join(html_parts)


# ---------------------------------------------------------------------------
# Plantillas HTML corporativas
# ---------------------------------------------------------------------------

SPECIES_TEMPLATE_ALIASES = {
    "porcino": {
        "label": "Porcino",
        "keywords": ["porc", "cerd", "lechon", "lechón", "marran", "ibérico", "iberico", "sow", "pig"],
        "photo": "foto_porcino.jpg",
        "template": "comercial_porcino.html",
    },
    "avicultura": {
        "label": "Avicultura",
        "keywords": ["avic", "gallina", "poll", "broiler", "pavo", "ponedora", "aves", "ave"],
        "photo": "foto_avicultura.jpg",
        "template": "comercial_avicultura.html",
    },
    "conejos": {
        "label": "Conejos",
        "keywords": ["conejo", "conej", "cuní", "cuni"],
        "photo": "foto_conejos.jpg",
        "template": "comercial_conejos.html",
    },
    "ovino": {
        "label": "Ovino",
        "keywords": ["ovino", "corder", "oveja", "borrego"],
        "photo": "foto_ovino.jpg",
        "template": "comercial_ovino.html",
    },
    "caprino": {
        "label": "Caprino",
        "keywords": ["caprino", "cabrit", "cabra", "chivo"],
        "photo": "foto_caprino.jpg",
        "template": "comercial_caprino.html",
    },
    "vacuno_leche": {
        "label": "Vacuno leche",
        "keywords": ["vacuno leche", "vacas leche", "vaca leche", "lactacion", "lactación", "lechera", "milk", "dairy"],
        "photo": "foto_vacuno_leche.jpg",
        "template": "comercial_vacuno_leche.html",
    },
    "vacuno_carne": {
        "label": "Vacuno carne",
        "keywords": ["vacuno carne", "ternero", "cebadero", "cebo vacuno", "nodriza", "beef", "carne"],
        "photo": "foto_vacuno_carne.jpg",
        "template": "comercial_vacuno_carne.html",
    },
    "caballos": {
        "label": "Caballos",
        "keywords": ["caballo", "caballos", "equino", "yegua", "potro", "horse"],
        "photo": "foto_caballos.jpg",
        "template": "comercial_caballos.html",
    },
}

BASE_TEMPLATE_BY_DOC_TYPE = {
    "Etiqueta": "etiqueta.html",
    "FT Calidad-Operaciones": "ft_calidad_operaciones.html",
    "FT Comercial": "comercial_base.html",
    "FT Especificaciones": "ft_especificaciones.html",
}

TEMPLATE_DISPLAY_NAMES = {
    "auto": "Automática recomendada",
    "nanta_corporativa.html": "Corporativa genérica NANTA",
    "etiqueta.html": "Etiqueta",
    "ft_calidad_operaciones.html": "FT Calidad-Operaciones",
    "ft_especificaciones.html": "FT Especificaciones",
    "comercial_base.html": "Comercial base",
    "comercial_porcino.html": "Comercial porcino",
    "comercial_avicultura.html": "Comercial avicultura",
    "comercial_conejos.html": "Comercial conejos",
    "comercial_ovino.html": "Comercial ovino",
    "comercial_caprino.html": "Comercial caprino",
    "comercial_vacuno_leche.html": "Comercial vacuno leche",
    "comercial_vacuno_carne.html": "Comercial vacuno carne",
    "comercial_caballos.html": "Comercial caballos",
}

DEFAULT_NANTA_HTML_TEMPLATE = """<!doctype html>
<html lang="es">
<head>
<meta charset="utf-8">
<title>{{TITLE}}</title>
<style>
  @page { size: A4; margin: 12mm; }
  * { box-sizing: border-box; }
  body { font-family: Arial, Helvetica, sans-serif; color: #2b2d38; margin: 0; background: #f3f4f7; }
  .page { width: 100%; max-width: 980px; margin: 0 auto; background:#fff; min-height: 1120px; position:relative; overflow:hidden; }
  .hero { position: relative; height: 255px; background: #eceff4; overflow:hidden; }
  .hero-photo { position:absolute; inset:0; width:100%; height:100%; object-fit:cover; }
  .hero::after { content:""; position:absolute; inset:0; background: linear-gradient(90deg, rgba(15,17,24,.76), rgba(15,17,24,.18)); }
  .logo-box { position:absolute; z-index:3; top:20px; left:24px; background:#fff; border-radius:16px; padding:13px 18px; box-shadow:0 8px 24px rgba(0,0,0,.16); }
  .logo { max-width: 185px; max-height: 80px; object-fit: contain; display:block; }
  .title-box { position:absolute; z-index:3; left:26px; bottom:24px; right:24px; color:#fff; }
  .kicker { font-weight:800; text-transform: uppercase; letter-spacing: .10em; font-size: 13px; }
  h1 { margin: 9px 0 10px; font-size: 36px; line-height: 1.04; }
  .meta { display: flex; flex-wrap: wrap; gap: 8px; }
  .pill { border-radius: 999px; padding: 7px 12px; font-size: 12px; background: rgba(255,255,255,.92); color:#222733; font-weight:700; }
  .solapa { width: 100%; height: 20px; object-fit: cover; display:block; }
  .content { padding: 20px 24px 24px; }
  .grid { display: grid; grid-template-columns: 1fr 1fr; gap: 14px; }
  .card { border: 1px solid #e0e1e8; border-radius: 14px; padding: 14px 16px; background: #fff; box-shadow: 0 2px 8px rgba(0,0,0,.05); break-inside: avoid; }
  .card h2 { color: #ef2533; margin: 0 0 9px; font-size: 14px; text-transform: uppercase; letter-spacing: .04em; border-bottom: 1px solid #f0b2ba; padding-bottom: 6px; }
  p { font-size: 13px; line-height: 1.40; margin: 0 0 7px; }
  table { width: 100%; border-collapse: collapse; font-size: 11.3px; margin-top: 6px; }
  th { text-align: left; background: #ef2533; color: #fff; padding: 5px; border: 1px solid #ef2533; }
  td { padding: 5px; border: 1px solid #d9dce5; vertical-align: top; }
  tr:nth-child(even) td { background: #fafafa; }
  .footer { margin-top: 18px; border-top: 1px solid #e0e1e8; padding-top: 8px; font-size: 10px; color: #6b6f7a; display: flex; justify-content: space-between; }
</style>
</head>
<body>
<div class="page">
  <header class="hero">
    <img class="hero-photo" src="{{SPECIES_PHOTO_URI}}" alt="">
    <div class="logo-box"><img class="logo" src="{{LOGO_URI}}" alt="NANTA"></div>
    <div class="title-box">
      <div class="kicker">{{DOC_TYPE}}</div>
      <h1>{{PRODUCT}}</h1>
      <div class="meta">
        <span class="pill">{{ESPECIE}}</span>
        <span class="pill">{{SUBESPECIE}}</span>
        <span class="pill">{{LIFESTAGE}}</span>
      </div>
    </div>
  </header>
  <img class="solapa" src="{{SOLAPA_URI}}" alt="">
  <main class="content">
    <section class="grid">{{SECTIONS}}</section>
    <div class="footer"><span>Ficha generada automáticamente. Revisar antes de uso definitivo.</span><span>{{DATE}}</span></div>
  </main>
</div>
</body>
</html>"""


def species_key_from_fields(fields: Dict[str, str]) -> str:
    haystack = " ".join(
        clean_text(fields.get(k, ""))
        for k in ["Especie", "Subespecie", "Lifestage", "Animales de destino", "Tipo de pienso", "Nombre comercial", "Producto"]
    )
    haystack_norm = strip_accents(haystack).lower()

    # Prioridades específicas para no confundir vacuno leche/carne con términos genéricos.
    if any(k in haystack_norm for k in ["vacuno leche", "vacas leche", "vaca leche", "lechera", "dairy"]):
        return "vacuno_leche"
    if any(k in haystack_norm for k in ["vacuno carne", "ternero", "cebo vacuno", "beef"]):
        return "vacuno_carne"

    for key, meta in SPECIES_TEMPLATE_ALIASES.items():
        if any(strip_accents(keyword).lower() in haystack_norm for keyword in meta["keywords"]):
            return key
    return "generico"


def species_label(species_key: str) -> str:
    return SPECIES_TEMPLATE_ALIASES.get(species_key, {}).get("label", "NANTA")


def default_species_cover_uri(species_key: str, label: str) -> str:
    palette = {
        "porcino": ("#f08aa0", "#712638"),
        "avicultura": ("#f7cc57", "#5b3e00"),
        "conejos": ("#d8dfe8", "#3d4958"),
        "ovino": ("#b8d7a0", "#315c2c"),
        "caprino": ("#c6d3ac", "#46572e"),
        "vacuno_leche": ("#b8dff1", "#1e5a72"),
        "vacuno_carne": ("#d9a06b", "#633717"),
        "caballos": ("#c79a75", "#54321f"),
        "generico": ("#ef2533", "#2b2d38"),
    }
    c1, c2 = palette.get(species_key, palette["generico"])
    title = html_lib.escape(label.upper())
    svg = f"""
    <svg xmlns='http://www.w3.org/2000/svg' width='1600' height='650' viewBox='0 0 1600 650'>
      <defs>
        <linearGradient id='g' x1='0' x2='1' y1='0' y2='1'>
          <stop offset='0%' stop-color='{c1}'/>
          <stop offset='100%' stop-color='{c2}'/>
        </linearGradient>
        <radialGradient id='r' cx='76%' cy='26%' r='64%'>
          <stop offset='0%' stop-color='rgba(255,255,255,.42)'/>
          <stop offset='100%' stop-color='rgba(255,255,255,0)'/>
        </radialGradient>
      </defs>
      <rect width='1600' height='650' fill='url(#g)'/>
      <rect width='1600' height='650' fill='url(#r)'/>
      <circle cx='1260' cy='150' r='210' fill='rgba(255,255,255,.18)'/>
      <circle cx='1380' cy='450' r='330' fill='rgba(255,255,255,.10)'/>
      <path d='M1040 430 C1120 360 1300 340 1430 420 C1310 500 1160 515 1040 430Z' fill='rgba(255,255,255,.18)'/>
      <text x='78' y='560' fill='rgba(255,255,255,.22)' font-family='Arial, Helvetica, sans-serif' font-size='118' font-weight='900'>{title}</text>
    </svg>
    """
    encoded = base64.b64encode(svg.encode("utf-8")).decode("ascii")
    return f"data:image/svg+xml;base64,{encoded}"


def species_photo_uri(fields: Dict[str, str]) -> Tuple[str, str, str]:
    species_key = species_key_from_fields(fields)
    label = species_label(species_key)
    file_name = SPECIES_TEMPLATE_ALIASES.get(species_key, {}).get("photo", "foto_generica.jpg")
    for suffix in [".jpg", ".jpeg", ".png", ".webp"]:
        path = ASSETS_DIR / Path(file_name).with_suffix(suffix).name
        uri = asset_to_data_uri(path)
        if uri:
            return uri, species_key, path.name
    return default_species_cover_uri(species_key, label), species_key, "fallback_svg"


def recommended_template_name(doc_type: str, fields: Dict[str, str]) -> str:
    if doc_type == "FT Comercial":
        species_key = species_key_from_fields(fields)
        if species_key in SPECIES_TEMPLATE_ALIASES:
            return SPECIES_TEMPLATE_ALIASES[species_key]["template"]
        return "comercial_base.html"
    return BASE_TEMPLATE_BY_DOC_TYPE.get(doc_type, "nanta_corporativa.html")


def list_template_files() -> List[str]:
    if not TEMPLATES_DIR.exists():
        return []
    return sorted(path.name for path in TEMPLATES_DIR.glob("*.html") if path.is_file())


def template_options_for_doc(doc_type: str, fields: Dict[str, str]) -> List[str]:
    files = list_template_files()
    recommended = recommended_template_name(doc_type, fields)
    options = ["auto"]

    ordered = []
    if recommended in files:
        ordered.append(recommended)
    base = BASE_TEMPLATE_BY_DOC_TYPE.get(doc_type)
    if base and base in files and base not in ordered:
        ordered.append(base)

    if doc_type == "FT Comercial":
        for name in [
            "comercial_base.html",
            "comercial_porcino.html",
            "comercial_avicultura.html",
            "comercial_conejos.html",
            "comercial_ovino.html",
            "comercial_caprino.html",
            "comercial_vacuno_leche.html",
            "comercial_vacuno_carne.html",
            "comercial_caballos.html",
        ]:
            if name in files and name not in ordered:
                ordered.append(name)

    for name in ["etiqueta.html", "ft_calidad_operaciones.html", "ft_especificaciones.html", "nanta_corporativa.html"]:
        if name in files and name not in ordered:
            ordered.append(name)

    for name in files:
        if name not in ordered:
            ordered.append(name)

    return options + ordered


def resolve_template_path(selected_template: str, doc_type: str, fields: Dict[str, str]) -> Tuple[Optional[Path], str]:
    name = recommended_template_name(doc_type, fields) if selected_template == "auto" else selected_template
    path = TEMPLATES_DIR / name
    if path.exists():
        return path, name

    fallback_name = BASE_TEMPLATE_BY_DOC_TYPE.get(doc_type, "nanta_corporativa.html")
    fallback_path = TEMPLATES_DIR / fallback_name
    if fallback_path.exists():
        return fallback_path, fallback_name

    generic_path = TEMPLATES_DIR / "nanta_corporativa.html"
    if generic_path.exists():
        return generic_path, "nanta_corporativa.html"

    return None, "plantilla_interna"


def html_card(field: str, value: str, css_class: str = "card") -> str:
    if not clean_text(value):
        return ""
    return f"""
    <section class="{css_class}">
      <h2>{html_lib.escape(field)}</h2>
      {section_to_html(value)}
    </section>
    """


def template_status_sidebar() -> None:
    with st.sidebar.expander("Plantillas y activos", expanded=False):
        files = list_template_files()
        if files:
            st.success(f"Plantillas HTML: {len(files)}")
            st.caption(", ".join(files[:8]) + ("…" if len(files) > 8 else ""))
        else:
            st.warning("No se han encontrado plantillas en /templates.")
        st.write("Logo:", "OK" if (ASSETS_DIR / "Logo1 Nanta.jpg").exists() else "No encontrado")
        st.write("Solapa:", "OK" if (ASSETS_DIR / "Solapa rosa.jpg").exists() else "No encontrada")
        photos = [meta["photo"] for meta in SPECIES_TEMPLATE_ALIASES.values() if (ASSETS_DIR / meta["photo"]).exists()]
        st.write("Fotos de especie:", f"{len(photos)} / {len(SPECIES_TEMPLATE_ALIASES)}")


def replace_template_markers(template: str, replacements: Dict[str, str]) -> str:
    rendered = template
    expanded: Dict[str, str] = {}
    for key, value in replacements.items():
        clean_key = clean_text(key)
        if not clean_key:
            continue
        value = "" if value is None else str(value)
        expanded[clean_key] = value
        expanded[clean_key.upper()] = value
        expanded[clean_key.lower()] = value
        expanded[norm_key(clean_key)] = value
        snake = re.sub(r"[^A-Za-z0-9]+", "_", strip_accents(clean_key)).strip("_").lower()
        expanded[snake] = value
        expanded[snake.upper()] = value

    for key in sorted(expanded, key=len, reverse=True):
        rendered = rendered.replace("{{" + key + "}}", expanded[key])
        rendered = rendered.replace("{{ " + key + " }}", expanded[key])
    return rendered


def render_nanta_html(
    doc_type: str,
    title: str,
    flat_rows: List[Dict[str, Any]],
    text: str,
    selected_template: str = "auto",
) -> Tuple[str, str, str]:
    fields = flat_rows_to_field_map(flat_rows)
    logo_uri = asset_to_data_uri(ASSETS_DIR / "Logo1 Nanta.jpg")
    solapa_uri = asset_to_data_uri(ASSETS_DIR / "Solapa rosa.jpg")
    cover_uri, species_key, cover_file = species_photo_uri(fields)

    product = fields.get("Nombre comercial") or fields.get("Producto") or title
    especie = fields.get("Especie", "")
    subespecie = fields.get("Subespecie", "")
    lifestage = fields.get("Lifestage", "")

    preferred_sections = [
        "Definición / Posicionamiento",
        "Características",
        "Beneficios",
        "Modo de empleo",
        "Precauciones de uso",
        "Animales de destino",
        "Nutrientes seleccionados",
        "Límites de nutrientes",
        "Límites de ingredientes",
    ]
    cards = []
    used_fields = {"Título", "Producto"}
    for field in preferred_sections:
        value = fields.get(field, "")
        if value:
            cards.append(html_card(field, value))
            used_fields.add(field)

    for row in flat_rows or []:
        field = clean_text(row.get("Campo", ""))
        value = clean_text(row.get("Valor", ""))
        if not field or not value or field in used_fields:
            continue
        cards.append(html_card(field, value))
        used_fields.add(field)

    if not cards:
        cards.append(f"<section class='card'><h2>Contenido</h2>{section_to_html(text)}</section>")

    template_path, template_name = resolve_template_path(selected_template, doc_type, fields)
    if template_path is not None:
        template = template_path.read_text(encoding="utf-8")
    else:
        template = DEFAULT_NANTA_HTML_TEMPLATE

    nutrientes_html = section_to_html(fields.get("Nutrientes seleccionados", ""))
    limites_nutrientes_html = section_to_html(fields.get("Límites de nutrientes", ""))
    limites_ingredientes_html = section_to_html(fields.get("Límites de ingredientes", ""))

    replacements = {
        "LOGO_URI": logo_uri,
        "SOLAPA_URI": solapa_uri,
        "SPECIES_PHOTO_URI": cover_uri,
        "PHOTO_URI": cover_uri,
        "DOC_TYPE": html_lib.escape(doc_type),
        "TITLE": html_lib.escape(title),
        "PRODUCT": html_lib.escape(product),
        "ESPECIE": html_lib.escape(especie),
        "SUBESPECIE": html_lib.escape(subespecie),
        "LIFESTAGE": html_lib.escape(lifestage),
        "SPECIES_KEY": html_lib.escape(species_key),
        "SPECIES_LABEL": html_lib.escape(species_label(species_key)),
        "SECTIONS": "\n".join(cards),
        "FIELDS_HTML": "\n".join(cards),
        "DATE": date.today().isoformat(),
        "BENEFICIOS": section_to_html(fields.get("Beneficios", "")),
        "BENEFICIO": section_to_html(fields.get("Beneficios", "")),
        "MODO_EMPLEO": section_to_html(fields.get("Modo de empleo", "")),
        "PRECAUCIONES_DE_USO": section_to_html(fields.get("Precauciones de uso", "")),
        "CARACTERISTICAS": section_to_html(fields.get("Características", "")),
        "DEFINICION_POSICIONAMIENTO": section_to_html(fields.get("Definición / Posicionamiento", "")),
        "NUTRIENTES_HTML": nutrientes_html,
        "LIMITES_NUTRIENTES_HTML": limites_nutrientes_html,
        "LIMITES_INGREDIENTES_HTML": limites_ingredientes_html,
        "NUTRIENTES_TEXTO": html_lib.escape(fields.get("Nutrientes seleccionados", "")).replace("\n", "<br>"),
        "LIMITES_NUTRIENTES_TEXTO": html_lib.escape(fields.get("Límites de nutrientes", "")).replace("\n", "<br>"),
        "LIMITES_INGREDIENTES_TEXTO": html_lib.escape(fields.get("Límites de ingredientes", "")).replace("\n", "<br>"),
        "TEMPLATE_NAME": html_lib.escape(template_name),
        "COVER_FILE": html_lib.escape(cover_file),
    }

    # Campos de usuario también como marcadores: {{Nombre comercial}}, {{nombre_comercial}}, etc.
    for field, value in fields.items():
        replacements[field] = html_lib.escape(clean_text(value)).replace("\n", "<br>")

    rendered = replace_template_markers(template, replacements)
    return rendered, template_name, cover_file


def build_canva_records(batch_rows: List[Dict[str, Any]]) -> pd.DataFrame:
    groups: Dict[Tuple[str, str], Dict[str, str]] = {}
    for row in batch_rows or []:
        doc_type = clean_text(row.get("Tipo documento", ""))
        product = clean_text(row.get("Producto", ""))
        field = clean_text(row.get("Campo", ""))
        value = clean_text(row.get("Valor", ""))
        if not doc_type and not product:
            continue
        key = (doc_type, product)
        groups.setdefault(key, {"tipo_documento": doc_type, "producto": product})
        if field:
            groups[key][field] = value

    records = []
    for fields in groups.values():
        record = {
            "tipo_documento": fields.get("tipo_documento", ""),
            "producto": fields.get("Nombre comercial") or fields.get("producto", ""),
            "especie": fields.get("Especie", ""),
            "subespecie": fields.get("Subespecie", ""),
            "lifestage": fields.get("Lifestage", ""),
            "tipo_pienso": fields.get("Tipo de pienso", ""),
            "animales_destino": fields.get("Animales de destino", ""),
            "definicion_posicionamiento": fields.get("Definición / Posicionamiento", ""),
            "definicion_corta": shorten_for_layout(fields.get("Definición / Posicionamiento", ""), 260),
            "caracteristicas": fields.get("Características", ""),
            "caracteristicas_cortas": shorten_for_layout(fields.get("Características", ""), 300),
            "foco_beneficio": fields.get("Foco-beneficio", ""),
            "beneficios": fields.get("Beneficios", ""),
            "beneficio_corto": shorten_for_layout(fields.get("Beneficios", ""), 320),
            "modo_empleo": fields.get("Modo de empleo", ""),
            "modo_empleo_corto": shorten_for_layout(fields.get("Modo de empleo", ""), 300),
            "precauciones": fields.get("Precauciones de uso", ""),
            "precauciones_cortas": shorten_for_layout(fields.get("Precauciones de uso", ""), 260),
            "nutrientes_texto": fields.get("Nutrientes seleccionados", ""),
            "limites_nutrientes_texto": fields.get("Límites de nutrientes", ""),
            "limites_ingredientes_texto": fields.get("Límites de ingredientes", ""),
        }
        records.append(record)

    return pd.DataFrame(records)



# ---------------------------------------------------------------------------
# Lectura de ficheros de entrada
# ---------------------------------------------------------------------------

def read_excel_raw_sheets(file_bytes: bytes, suffix: str) -> Dict[str, pd.DataFrame]:
    engine = None
    suffix = suffix.lower()
    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        engine = "openpyxl"
    elif suffix == ".xls":
        engine = "xlrd"
    elif suffix == ".ods":
        engine = "odf"
    elif suffix == ".xlsb":
        engine = "pyxlsb"

    bio = io.BytesIO(file_bytes)
    xls = pd.ExcelFile(bio, engine=engine)
    sheets: Dict[str, pd.DataFrame] = {}
    for sheet_name in xls.sheet_names:
        bio2 = io.BytesIO(file_bytes)
        try:
            sheets[sheet_name] = pd.read_excel(
                bio2,
                sheet_name=sheet_name,
                header=None,
                dtype=object,
                engine=engine,
            )
        except Exception:
            # Fall back without explicit engine, useful in some cloud deployments.
            bio3 = io.BytesIO(file_bytes)
            sheets[sheet_name] = pd.read_excel(
                bio3,
                sheet_name=sheet_name,
                header=None,
                dtype=object,
            )
    return sheets


def read_csv_or_text(file_bytes: bytes, suffix: str) -> Tuple[Dict[str, pd.DataFrame], str]:
    raw_text = decode_bytes(file_bytes)
    sheets: Dict[str, pd.DataFrame] = {}

    if suffix.lower() in {".csv", ".tsv", ".txt", ".dat", ".prn"}:
        sep = "\t" if suffix.lower() == ".tsv" else None
        try:
            df = pd.read_csv(
                io.StringIO(raw_text),
                sep=sep,
                engine="python",
                dtype=object,
                on_bad_lines="skip",
            )
            sheets["Datos"] = df
        except Exception:
            # A pure text formulation output is still valid.
            lines = raw_text.splitlines()
            sheets["Texto"] = pd.DataFrame({"Texto": lines})
    return sheets, raw_text


def read_pdf_text(file_bytes: bytes) -> str:
    if PdfReader is None:
        st.warning("Para leer PDF instale pypdf. El fichero PDF no se ha podido leer.")
        return ""
    reader = PdfReader(io.BytesIO(file_bytes))
    parts: List[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts)


def read_docx_text(file_bytes: bytes) -> str:
    if Document is None:
        st.warning("Para leer DOCX instale python-docx. El fichero DOCX no se ha podido leer.")
        return ""
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if clean_text(p.text)]
    for table in doc.tables:
        for row in table.rows:
            vals = [clean_text(cell.text) for cell in row.cells]
            vals = [v for v in vals if v]
            if vals:
                parts.append(" | ".join(vals))
    return "\n".join(parts)


def read_json_file(file_bytes: bytes) -> Tuple[Dict[str, pd.DataFrame], str]:
    raw_text = decode_bytes(file_bytes)
    try:
        obj = json.loads(raw_text)
    except Exception:
        return {"Texto": pd.DataFrame({"Texto": raw_text.splitlines()})}, raw_text

    if isinstance(obj, list):
        df = pd.DataFrame(obj)
    elif isinstance(obj, dict):
        # Try to flatten common structures. If not possible, store key/value.
        if any(isinstance(v, list) for v in obj.values()):
            dfs = {}
            for k, v in obj.items():
                if isinstance(v, list):
                    dfs[str(k)] = pd.DataFrame(v)
            if dfs:
                return dfs, raw_text
        df = pd.json_normalize(obj)
    else:
        df = pd.DataFrame({"Valor": [obj]})
    return {"Datos": df}, raw_text


def promote_header(raw: pd.DataFrame, required_terms: Optional[List[str]] = None) -> pd.DataFrame:
    """Promote the most probable header row to DataFrame columns."""
    if raw.empty:
        return raw.copy()

    required_terms = required_terms or []
    best_idx = 0
    best_score = -1

    max_rows = min(len(raw), 20)
    for i in range(max_rows):
        row_values = [clean_text(v) for v in raw.iloc[i].tolist()]
        row_keys = [norm_key(v) for v in row_values if v]
        if not row_keys:
            continue
        score = len(row_keys)
        for term in required_terms:
            t = norm_key(term)
            if any(t in k for k in row_keys):
                score += 10
        if score > best_score:
            best_score = score
            best_idx = i

    header = [clean_text(v) for v in raw.iloc[best_idx].tolist()]
    cols = []
    used = {}
    for idx, col in enumerate(header):
        col = col or f"Columna_{idx + 1}"
        if col in used:
            used[col] += 1
            col = f"{col}_{used[col]}"
        else:
            used[col] = 1
        cols.append(col)
    df = raw.iloc[best_idx + 1:].copy()
    df.columns = cols
    df = df.dropna(how="all")
    return df.reset_index(drop=True)


# ---------------------------------------------------------------------------
# Lectura y normalización de "Etiquetas y Beneficios"
# ---------------------------------------------------------------------------

PARAM_ALIASES = {
    "especie": "Especie",
    "subespecie": "Subespecie",
    "animalesdedestino": "Animales de destino",
    "animaldestino": "Animales de destino",
    "destino": "Animales de destino",
    "opcion": "Opción",
    "tipo": "Tipo de pienso",
    "tipodepienso": "Tipo de pienso",
    "focobeneficio": "Foco-beneficio",
    "focobeneficioasociado": "Foco-beneficio",
    "modo": "Modo de empleo",
    "mododeempleo": "Modo de empleo",
    "mododeempleouso": "Modo de empleo",
    "precauciones": "Precauciones de uso",
    "precaucionesdeuso": "Precauciones de uso",
    "recomendaciones": "Precauciones de uso",
    "estadoproductivo": "Lifestage",
    "lifestage": "Lifestage",
    "texto": "Texto beneficio",
    "textobeneficio": "Texto beneficio",
    "beneficio": "Texto beneficio",
}


def standardize_columns(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    new_cols = {}
    for col in out.columns:
        key = norm_key(col)
        if key in PARAM_ALIASES:
            new_cols[col] = PARAM_ALIASES[key]
        else:
            # Heuristic matching for combined names.
            mapped = None
            for alias_key, canonical in PARAM_ALIASES.items():
                if alias_key and alias_key in key:
                    mapped = canonical
                    break
            new_cols[col] = mapped or clean_text(col)
    out = out.rename(columns=new_cols)

    # Drop completely empty columns and rows.
    out = out.dropna(axis=1, how="all")
    out = out.dropna(axis=0, how="all")
    for col in out.columns:
        out[col] = out[col].map(clean_text)
    return out.reset_index(drop=True)


def load_parameter_file(uploaded_file) -> Tuple[pd.DataFrame, pd.DataFrame]:
    if uploaded_file is None:
        return pd.DataFrame(), pd.DataFrame()

    if isinstance(uploaded_file, (str, Path)):
        path = Path(uploaded_file)
        file_bytes = path.read_bytes()
        suffix = path.suffix.lower()
    else:
        file_bytes = uploaded_file.getvalue()
        suffix = Path(uploaded_file.name).suffix.lower()

    if suffix not in {".xlsx", ".xlsm", ".xls", ".xlsb", ".ods"}:
        st.error("El fichero de textos parametrizados debe ser Excel/ODS con hojas Etiquetas y Beneficios.")
        return pd.DataFrame(), pd.DataFrame()

    raw_sheets = read_excel_raw_sheets(file_bytes, suffix)
    etiquetas = pd.DataFrame()
    beneficios = pd.DataFrame()

    for sheet_name, raw_df in raw_sheets.items():
        sheet_key = norm_key(sheet_name)
        if "etiqueta" in sheet_key:
            df = promote_header(raw_df, ["Especie", "Subespecie", "Modo de empleo", "Precauciones"])
            etiquetas = standardize_columns(df)
        elif "beneficio" in sheet_key:
            df = promote_header(raw_df, ["Especie", "Subespecie", "Estado productivo", "Texto beneficio"])
            beneficios = standardize_columns(df)

    # Fallback if sheets are not named as expected.
    if etiquetas.empty or beneficios.empty:
        for sheet_name, raw_df in raw_sheets.items():
            df = standardize_columns(promote_header(raw_df, ["Especie", "Subespecie"]))
            cols = {norm_key(c) for c in df.columns}
            if etiquetas.empty and ("mododeempleo" in cols or any("precauc" in c for c in cols)):
                etiquetas = df
            if beneficios.empty and ("textobeneficio" in cols or "lifestage" in cols):
                beneficios = df

    return etiquetas, beneficios


def filter_df(df: pd.DataFrame, **filters: str) -> pd.DataFrame:
    if df.empty:
        return df
    out = df.copy()
    for col, value in filters.items():
        if not value or col not in out.columns:
            continue
        out = out[out[col].map(norm_key) == norm_key(value)]
    return out


def get_unique(df: pd.DataFrame, column: str) -> List[str]:
    if df.empty or column not in df.columns:
        return []
    return unique_clean(df[column].tolist())


def option_label(row: pd.Series, columns: List[str], max_len: int = 120) -> str:
    parts = []
    for col in columns:
        if col in row.index and clean_text(row[col]):
            parts.append(clean_text(row[col]))
    text = " | ".join(parts)
    if len(text) > max_len:
        text = text[: max_len - 1] + "…"
    return text or "Opción sin texto"


# ---------------------------------------------------------------------------
# Parser de formulaciones Multi-Mix en texto
# ---------------------------------------------------------------------------

def split_multimix_blocks(raw_text: str) -> List[str]:
    if not raw_text:
        return []
    matches = list(re.finditer(r"(?im)Specification\s*:", raw_text))
    if not matches:
        return []
    blocks = []
    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(raw_text)
        blocks.append(raw_text[start:end].strip())
    return blocks


def parse_spec_line(block: str) -> Tuple[str, str, str]:
    first_lines = "\n".join(block.splitlines()[:5])
    m = re.search(
        r"Specification\s*:\s*(?P<spec>\S+)\s+(?P<name>.*?)\s*:\s*Cost/tonne\s*:\s*(?P<cost>[-0-9.,]+)",
        first_lines,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if not m:
        # Fallback: first non-empty line.
        line = next((clean_text(x) for x in block.splitlines() if clean_text(x)), "Producto sin nombre")
        return "", line[:80], ""
    spec = clean_text(m.group("spec"))
    name = clean_text(m.group("name"))
    cost = clean_text(m.group("cost"))
    name = re.sub(r"\s+", " ", name).strip(" :")
    return spec, name, cost


def parse_multimix_ingredients(block: str) -> pd.DataFrame:
    lines = block.splitlines()
    start = None
    for i, line in enumerate(lines):
        if "INCLUDED RAW MATERIALS" in line.upper():
            start = i + 1
            break
    if start is None:
        return pd.DataFrame()

    rows: List[Dict[str, Any]] = []
    fallback_re = re.compile(
        r"^\s*(?P<line_code>[A-Z]?\d+)\s+"
        r"(?P<material>.+?)\s+"
        r"(?P<cost>-?\d+(?:[\.,]\d+)?)\s+"
        r"(?P<pct>-?\d+(?:[\.,]\d+)?)\s+"
        r"(?P<kilos>-?\d+(?:[\.,]\d+)?)\s+"
        r"(?P<tonnes>-?\d+(?:[\.,]\d+)?)"
        r"(?:\s+(?P<limit>MAX|MIN|min|max|FIJO|Fijo|fix|FIX))?"
        r"(?:\s+(?P<minimum>\.|-?\d+(?:[\.,]\d+)?))?"
        r"(?:\s+(?P<maximum>\.|-?\d+(?:[\.,]\d+)?))?",
        flags=re.IGNORECASE,
    )

    for line in lines[start:]:
        if "ANALYSIS" in line.upper():
            break
        raw_line = str(line).rstrip()
        text = clean_text(line)
        if not text or set(text) <= {"-"}:
            continue
        if not re.search(r"^\s*[A-Z]?\d+", raw_line):
            continue

        parts = re.split(r"\s{2,}", raw_line.strip())
        parsed = None

        if len(parts) >= 5:
            first = parts[0]
            m = re.match(r"(?P<line_code>\S+)\s+(?P<material>.+)", first)
            if m:
                after = parts[1:]
                idx = 4
                limit = ""
                if len(after) > idx and norm_key(after[idx]) in {"min", "max", "fijo", "fix"}:
                    limit = after[idx]
                    idx += 1
                parsed = {
                    "line_code": m.group("line_code"),
                    "material": m.group("material").strip(),
                    "cost": after[0] if len(after) > 0 else "",
                    "pct": after[1] if len(after) > 1 else "",
                    "kilos": after[2] if len(after) > 2 else "",
                    "tonnes": after[3] if len(after) > 3 else "",
                    "limit": limit,
                    "minimum": after[idx] if len(after) > idx else "",
                    "maximum": after[idx + 1] if len(after) > idx + 1 else "",
                }

        if parsed is None:
            m = fallback_re.match(raw_line)
            if not m:
                continue
            parsed = m.groupdict()

        material_full = clean_text(parsed.get("material", ""))
        material_code = ""
        material_name = material_full
        m2 = re.match(r"(?P<code>\d+)\.(?P<name>.+)", material_full)
        if m2:
            material_code = m2.group("code")
            material_name = m2.group("name").strip()

        rows.append(
            {
                "Código línea": clean_text(parsed.get("line_code", "")),
                "Código materia prima": material_code,
                "Materia prima": material_name,
                "Coste medio": parse_float(parsed.get("cost", "")),
                "%": parse_float(parsed.get("pct", "")),
                "Kilos": parse_float(parsed.get("kilos", "")),
                "Toneladas": parse_float(parsed.get("tonnes", "")),
                "Restricción": clean_text(parsed.get("limit", "")),
                "Mínimo": clean_text(parsed.get("minimum", "")),
                "Máximo": clean_text(parsed.get("maximum", "")),
                "Línea original": text,
            }
        )

    return pd.DataFrame(rows)


def parse_multimix_analysis(block: str) -> pd.DataFrame:
    lines = block.splitlines()
    start = None
    for i, line in enumerate(lines):
        if re.search(r"\bANALYSIS\b", line, flags=re.IGNORECASE):
            start = i + 1
            break
    if start is None:
        return pd.DataFrame()

    rows: List[Dict[str, Any]] = []
    fallback_re = re.compile(
        r"^\s*(?P<name>\[[^\]]+\]|[A-Za-z0-9_+/.\-]+)\s+"
        r"(?P<level>\.|-?\d+(?:[\.,]\d+)?)"
        r"(?:\s+(?P<restriction>MAX|MIN|min|max|FIJO|Fijo|fix|FIX))?"
        r"(?:\s+(?P<minimum>\.|-?\d+(?:[\.,]\d+)?))?"
        r"(?:\s+(?P<maximum>\.|-?\d+(?:[\.,]\d+)?))?",
        flags=re.IGNORECASE,
    )

    for line in lines[start:]:
        raw_line = str(line).rstrip()
        text = clean_text(line)
        if not text or set(text) <= {"-"}:
            continue
        if "INCLUDED RAW MATERIALS" in text.upper() or "OPTIMIZATION" in text.upper():
            break

        parts = re.split(r"\s{2,}", raw_line.strip())
        parsed = None
        if len(parts) >= 2 and is_numeric_text(parts[1]):
            idx = 2
            restriction = ""
            if len(parts) > idx and norm_key(parts[idx]) in {"min", "max", "fijo", "fix"}:
                restriction = parts[idx]
                idx += 1
            parsed = {
                "name": parts[0],
                "level": parts[1],
                "restriction": restriction,
                "minimum": parts[idx] if len(parts) > idx else "",
                "maximum": parts[idx + 1] if len(parts) > idx + 1 else "",
            }
        else:
            m = fallback_re.match(raw_line)
            if m and is_numeric_text(m.group("level")):
                parsed = m.groupdict()

        if not parsed:
            continue

        rows.append(
            {
                "Nutriente": clean_text(parsed.get("name", "")),
                "Valor": parse_float(parsed.get("level", "")),
                "Restricción": clean_text(parsed.get("restriction", "")),
                "Mínimo": clean_text(parsed.get("minimum", "")),
                "Máximo": clean_text(parsed.get("maximum", "")),
                "Línea original": text,
            }
        )

    return pd.DataFrame(rows)


def parse_multimix_text(raw_text: str) -> Dict[str, Dict[str, Any]]:
    products: Dict[str, Dict[str, Any]] = {}
    for idx, block in enumerate(split_multimix_blocks(raw_text), start=1):
        spec, name, cost = parse_spec_line(block)
        if not name:
            name = f"Producto {idx}"
        ingredients = parse_multimix_ingredients(block)
        nutrients = parse_multimix_analysis(block)

        display = f"{spec} | {name}" if spec else name
        # Avoid duplicates.
        base_display = display
        n = 2
        while display in products:
            display = f"{base_display} ({n})"
            n += 1

        products[display] = {
            "product_name": name,
            "product_id": spec,
            "cost_per_tonne": cost,
            "source_type": "Multi-Mix texto",
            "row": {},
            "nutrients": nutrients,
            "ingredients": ingredients,
            "nutrient_limits": nutrients[
                nutrients[["Mínimo", "Máximo", "Restricción"]].fillna("").astype(str).agg("".join, axis=1).str.strip() != ""
            ].copy() if not nutrients.empty else pd.DataFrame(),
            "ingredient_limits": ingredients[
                ingredients[["Mínimo", "Máximo", "Restricción"]].fillna("").astype(str).agg("".join, axis=1).str.strip() != ""
            ].copy() if not ingredients.empty else pd.DataFrame(),
            "raw_text": block,
        }
    return products



# ---------------------------------------------------------------------------
# Parser de formulaciones Single-Mix en texto con bloques SP:
# ---------------------------------------------------------------------------

def clean_limit_value(value: Any) -> str:
    text = clean_text(value)
    return "" if text in {".", "-", "--"} else text


def split_singlemix_blocks(raw_text: str) -> List[str]:
    """Split Single-Mix reports by product blocks marked as ': SP: CODE NAME ...'."""
    if not raw_text:
        return []
    matches = list(re.finditer(r"(?im)^\s*:?\s*SP:\s+", raw_text))
    if not matches:
        return []
    blocks: List[str] = []
    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(raw_text)
        block = raw_text[start:end].strip()
        if block:
            blocks.append(block)
    return blocks


def parse_singlemix_header(block: str) -> Tuple[str, str, str]:
    header = "\n".join(block.splitlines()[:8])
    m = re.search(
        r"SP:\s*(?P<spec>\S+)\s+(?P<name>.*?)\s+"
        r"(?P<pct>-?\d+(?:[\.,]\d+)?)\s*%,\s*"
        r"(?P<kg>-?\d+(?:[\.,]\d+)?)\s*Kg.*?Optimal\s+cost:\s*(?P<cost>-?\d+(?:[\.,]\d+)?)",
        header,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if not m:
        line = next((clean_text(x) for x in block.splitlines() if "SP:" in x.upper()), "")
        m2 = re.search(r"SP:\s*(?P<spec>\S+)\s+(?P<name>.+)", line, flags=re.IGNORECASE)
        if m2:
            return clean_text(m2.group("spec")), re.sub(r"\s{2,}.*$", "", clean_text(m2.group("name"))).strip(), ""
        return "", "Producto sin nombre", ""
    spec = clean_text(m.group("spec"))
    name = re.sub(r"\s+", " ", clean_text(m.group("name"))).strip(" :")
    cost = clean_text(m.group("cost"))
    return spec, name, cost


def _material_code_name(material_full: str) -> Tuple[str, str]:
    material_full = clean_text(material_full)
    material_code = ""
    material_name = material_full
    m = re.match(r"(?P<code>\d+)\.(?P<name>.+)", material_full)
    if m:
        material_code = clean_text(m.group("code"))
        material_name = clean_text(m.group("name"))
    return material_code, material_name


def _parse_singlemix_material_line(line: str, section: str) -> Optional[Dict[str, Any]]:
    raw_line = str(line).rstrip()
    text = clean_text(raw_line)
    if not text or set(text) <= {"-"}:
        return None
    if not re.match(r"^\s*(?:[A-Z]?\d+|X\d+)\b", raw_line):
        return None

    parts = re.split(r"\s{2,}", raw_line.strip())
    if len(parts) < 4:
        return None

    # Some lines are split as ['X394', '12394765.NP-326...', '0.3', ...]
    # and others as ['5128 10079340.CEBADA...', '48.0', ...].
    line_code = ""
    material_full = ""
    offset = 1
    if re.fullmatch(r"[A-Z]?\d+|X\d+", parts[0], flags=re.IGNORECASE) and len(parts) >= 5 and not is_numeric_text(parts[1]):
        line_code = clean_text(parts[0])
        material_full = clean_text(parts[1])
        offset = 2
    else:
        m = re.match(r"(?P<line_code>[A-Z]?\d+|X\d+)\s+(?P<material>.+)", parts[0], flags=re.IGNORECASE)
        if not m:
            return None
        line_code = clean_text(m.group("line_code"))
        material_full = clean_text(m.group("material"))
        offset = 1

    material_code, material_name = _material_code_name(material_full)

    if section == "included":
        if len(parts) < offset + 3:
            return None
        pct = parts[offset]
        kg = parts[offset + 1]
        cost = parts[offset + 2]
        idx = offset + 3
    else:
        # Rejected raw materials: no actual %/kg. Columns start with minimum/maximum/cost.
        pct = "0"
        kg = ""
        cost = parts[offset + 2] if len(parts) > offset + 2 else ""
        idx = offset

    restriction = ""
    if len(parts) > idx and norm_key(parts[idx]) in {"min", "max", "fijo", "fix"}:
        restriction = clean_text(parts[idx])
        idx += 1

    minimum = clean_limit_value(parts[idx]) if len(parts) > idx else ""
    maximum = clean_limit_value(parts[idx + 1]) if len(parts) > idx + 1 else ""

    return {
        "Código línea": line_code,
        "Código materia prima": material_code,
        "Materia prima": material_name,
        "Coste medio": parse_float(cost),
        "%": parse_float(pct),
        "Kilos": parse_float(kg),
        "Toneladas": "",
        "Restricción": restriction,
        "Mínimo": minimum,
        "Máximo": maximum,
        "Origen": "Incluida" if section == "included" else "Rechazada",
        "Línea original": text,
    }


def parse_singlemix_materials(block: str, section: str = "included") -> pd.DataFrame:
    lines = block.splitlines()
    start = None
    stop_terms = []
    if section == "included":
        start_term = "INCLUDED RAW MATERIALS"
        stop_terms = ["REJECTED RAW MATERIALS", "NUTRIENT ANALYSIS", "OPTIMIZATION"]
    else:
        start_term = "REJECTED RAW MATERIALS"
        stop_terms = ["NUTRIENT ANALYSIS", "OPTIMIZATION"]

    for i, line in enumerate(lines):
        if start_term in line.upper():
            start = i + 1
            break
    if start is None:
        return pd.DataFrame()

    rows: List[Dict[str, Any]] = []
    for line in lines[start:]:
        upper = line.upper()
        if any(term in upper for term in stop_terms):
            break
        parsed = _parse_singlemix_material_line(line, section)
        if parsed:
            rows.append(parsed)
    return pd.DataFrame(rows)


def parse_singlemix_analysis(block: str) -> pd.DataFrame:
    lines = block.splitlines()
    start = None
    for i, line in enumerate(lines):
        if "NUTRIENT ANALYSIS" in line.upper():
            start = i + 1
            break
    if start is None:
        return pd.DataFrame()

    rows: List[Dict[str, Any]] = []
    for line in lines[start:]:
        raw_line = str(line).rstrip()
        text = clean_text(raw_line)
        if not text or set(text) <= {"-"}:
            continue
        upper = text.upper()
        if "SP:" in upper or "OPTIMIZATION" in upper or "INCLUDED RAW MATERIALS" in upper:
            break

        parts = re.split(r"\s{2,}", raw_line.strip())
        if len(parts) < 2:
            continue

        name = clean_text(parts[0])
        unit = ""
        if len(parts) >= 3 and not is_numeric_text(parts[1]) and is_numeric_text(parts[2]):
            unit = clean_text(parts[1])
            level_idx = 2
        elif is_numeric_text(parts[1]):
            level_idx = 1
        else:
            continue

        restriction = ""
        idx = level_idx + 1
        if len(parts) > idx and norm_key(parts[idx]) in {"min", "max", "fijo", "fix"}:
            restriction = clean_text(parts[idx])
            idx += 1

        minimum = clean_limit_value(parts[idx]) if len(parts) > idx else ""
        maximum = clean_limit_value(parts[idx + 1]) if len(parts) > idx + 1 else ""

        rows.append({
            "Nutriente": name,
            "Unidad": unit,
            "Valor": parse_float(parts[level_idx]),
            "Restricción": restriction,
            "Mínimo": minimum,
            "Máximo": maximum,
            "Línea original": text,
        })

    return pd.DataFrame(rows)


def parse_singlemix_text(raw_text: str) -> Dict[str, Dict[str, Any]]:
    products: Dict[str, Dict[str, Any]] = {}
    for idx, block in enumerate(split_singlemix_blocks(raw_text), start=1):
        spec, name, cost = parse_singlemix_header(block)
        if not name:
            name = f"Producto {idx}"

        included = parse_singlemix_materials(block, "included")
        rejected = parse_singlemix_materials(block, "rejected")
        nutrients = parse_singlemix_analysis(block)

        if included.empty and nutrients.empty:
            continue

        ingredient_limits_parts = []
        if not included.empty:
            ingredient_limits_parts.append(included)
        if not rejected.empty:
            ingredient_limits_parts.append(rejected)
        ingredient_limits_all = pd.concat(ingredient_limits_parts, ignore_index=True) if ingredient_limits_parts else pd.DataFrame()

        display = f"{spec} | {name}" if spec else name
        base_display = display
        n = 2
        while display in products:
            display = f"{base_display} ({n})"
            n += 1

        products[display] = {
            "product_name": name,
            "product_id": spec,
            "cost_per_tonne": cost,
            "source_type": "Single-Mix texto SP",
            "row": {},
            "nutrients": nutrients,
            "ingredients": included,
            "nutrient_limits": nutrients[
                nutrients[["Mínimo", "Máximo", "Restricción"]].fillna("").astype(str).agg("".join, axis=1).str.strip() != ""
            ].copy() if not nutrients.empty else pd.DataFrame(),
            "ingredient_limits": ingredient_limits_all[
                ingredient_limits_all[["Mínimo", "Máximo", "Restricción"]].fillna("").astype(str).agg("".join, axis=1).str.strip() != ""
            ].copy() if not ingredient_limits_all.empty else pd.DataFrame(),
            "raw_text": block,
        }
    return products


# ---------------------------------------------------------------------------
# Parser genérico de tablas de productos
# ---------------------------------------------------------------------------

def choose_product_column(df: pd.DataFrame) -> Optional[str]:
    if df.empty:
        return None
    priority_terms = [
        "producto",
        "nombrecomercial",
        "nombre",
        "pienso",
        "descripcion",
        "description",
        "product",
        "specification",
        "formula",
        "codigo",
    ]
    for col in df.columns:
        key = norm_key(col)
        if any(term in key for term in priority_terms):
            return col

    # Fallback: text column with a useful number of unique values.
    best_col = None
    best_score = 0
    for col in df.columns:
        values = df[col].dropna().map(clean_text)
        if values.empty:
            continue
        non_numeric = values[~values.map(is_numeric_text)]
        score = non_numeric.nunique()
        if score > best_score and score >= 2:
            best_score = score
            best_col = col
    return best_col


def extract_products_from_tables(raw_sheets: Dict[str, pd.DataFrame]) -> Dict[str, Dict[str, Any]]:
    products: Dict[str, Dict[str, Any]] = {}

    for sheet_name, raw_df in raw_sheets.items():
        df = promote_header(raw_df, ["Producto", "Nombre", "Pienso", "Specification"])
        df = df.dropna(axis=1, how="all").dropna(axis=0, how="all")
        if df.empty:
            continue
        product_col = choose_product_column(df)
        if not product_col:
            continue

        for row_idx, row in df.iterrows():
            product_name = clean_text(row.get(product_col, ""))
            if not product_name:
                continue
            row_dict = {clean_text(k): clean_text(v) for k, v in row.to_dict().items() if clean_text(v)}

            nutrients_rows = []
            for col in df.columns:
                col_key = norm_key(col)
                if col == product_col:
                    continue
                value = row.get(col)
                if is_numeric_text(value):
                    # Exclude fields that are clearly IDs/codes rather than nutrient values.
                    if any(x in col_key for x in ["codigo", "ean", "fecha", "version", "registro"]):
                        continue
                    nutrients_rows.append(
                        {
                            "Nutriente": clean_text(col),
                            "Valor": parse_float(value),
                            "Restricción": "",
                            "Mínimo": "",
                            "Máximo": "",
                            "Línea original": f"{clean_text(col)} {clean_text(value)}",
                        }
                    )

            ingredients_rows = []
            for col in df.columns:
                col_key = norm_key(col)
                if any(x in col_key for x in ["ingred", "materiaprima", "composicion", "rawmaterial"]):
                    text = clean_text(row.get(col, ""))
                    if text:
                        ingredients_rows.append({"Materia prima": clean_text(col), "%": "", "Línea original": text})

            display = f"{sheet_name} | {product_name}"
            base_display = display
            n = 2
            while display in products:
                display = f"{base_display} ({n})"
                n += 1

            products[display] = {
                "product_name": product_name,
                "product_id": row_dict.get("Código", row_dict.get("Codigo", "")),
                "cost_per_tonne": "",
                "source_type": f"Tabla: {sheet_name}",
                "row": row_dict,
                "nutrients": pd.DataFrame(nutrients_rows),
                "ingredients": pd.DataFrame(ingredients_rows),
                "nutrient_limits": pd.DataFrame(),
                "ingredient_limits": pd.DataFrame(),
                "raw_text": "",
            }

    return products


def load_formula_file(uploaded_file) -> Dict[str, Dict[str, Any]]:
    if uploaded_file is None:
        return {}

    file_bytes = uploaded_file.getvalue()
    suffix = Path(uploaded_file.name).suffix.lower()

    raw_text = ""
    raw_sheets: Dict[str, pd.DataFrame] = {}

    try:
        if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm", ".xls", ".xlsb", ".ods"}:
            raw_sheets = read_excel_raw_sheets(file_bytes, suffix)
            raw_text = "\n".join(df_to_raw_text(df) for df in raw_sheets.values())
        elif suffix in {".csv", ".tsv", ".txt", ".dat", ".prn"}:
            raw_sheets, raw_text = read_csv_or_text(file_bytes, suffix)
        elif suffix == ".json":
            raw_sheets, raw_text = read_json_file(file_bytes)
        elif suffix == ".pdf":
            raw_text = read_pdf_text(file_bytes)
            raw_sheets = {"Texto PDF": pd.DataFrame({"Texto": raw_text.splitlines()})}
        elif suffix == ".docx":
            raw_text = read_docx_text(file_bytes)
            raw_sheets = {"Texto DOCX": pd.DataFrame({"Texto": raw_text.splitlines()})}
        elif suffix == ".parquet":
            df = pd.read_parquet(io.BytesIO(file_bytes))
            raw_sheets = {"Datos": df}
            raw_text = df_to_raw_text(df)
        else:
            # Last attempt: interpret as text.
            raw_text = decode_bytes(file_bytes)
            raw_sheets = {"Texto": pd.DataFrame({"Texto": raw_text.splitlines()})}
    except Exception as exc:
        st.error(f"No se ha podido leer el fichero de formulación: {exc}")
        return {}

    products = parse_multimix_text(raw_text)
    if products:
        return products

    products = parse_singlemix_text(raw_text)
    if products:
        return products

    products = extract_products_from_tables(raw_sheets)
    return products


# ---------------------------------------------------------------------------
# Inferencia básica producto -> especie/subespecie/lifestage
# ---------------------------------------------------------------------------

def infer_defaults_from_product(
    product_name: str,
    etiquetas: pd.DataFrame,
    beneficios: pd.DataFrame,
) -> Dict[str, str]:
    text = norm_key(product_name)
    defaults = {"Especie": "", "Subespecie": "", "Lifestage": ""}

    # Keyword rules. They are intentionally conservative.
    if any(k in text for k in ["optibaby", "lechon", "porc", "pig"]):
        defaults["Especie"] = "PORCINO"
        defaults["Subespecie"] = "BLANCO"
        defaults["Lifestage"] = "LECHONES"
    elif any(k in text for k in ["gesticor", "corder", "cabrit", "ovino", "caprino"]):
        defaults["Especie"] = "PEQUEÑO RUMIANTE"
        defaults["Subespecie"] = "OVINO/CAPRINO CARNE"
        defaults["Lifestage"] = "CEBO"
    elif any(k in text for k in ["gallina", "ponedora", "puesta"]):
        defaults["Especie"] = "GALLINAS"
        defaults["Subespecie"] = "AC10 PUESTA"
    elif any(k in text for k in ["pollo", "broiler"]):
        defaults["Especie"] = "POLLOS"
        defaults["Subespecie"] = "AC20 POLLOS"
    elif any(k in text for k in ["pavo", "turkey"]):
        defaults["Especie"] = "PAVOS"
        defaults["Subespecie"] = "AC30 PAVOS"
    elif any(k in text for k in ["vacuno", "ternero", "lactacion", "rumiante"]):
        defaults["Especie"] = "RUMIANTES"

    available_species = set(get_unique(etiquetas, "Especie") + get_unique(beneficios, "Especie"))
    if defaults["Especie"] and defaults["Especie"] not in available_species:
        defaults["Especie"] = ""

    available_subspecies = set(get_unique(filter_df(etiquetas, Especie=defaults["Especie"]), "Subespecie") +
                               get_unique(filter_df(beneficios, Especie=defaults["Especie"]), "Subespecie"))
    if defaults["Subespecie"] and defaults["Subespecie"] not in available_subspecies:
        defaults["Subespecie"] = ""

    available_lifestages = set(get_unique(
        filter_df(beneficios, Especie=defaults["Especie"], Subespecie=defaults["Subespecie"]),
        "Lifestage",
    ))
    if defaults["Lifestage"] and defaults["Lifestage"] not in available_lifestages:
        # If a specific lifestage is unavailable, keep first available for the chosen subspecies.
        defaults["Lifestage"] = sorted(available_lifestages)[0] if available_lifestages else ""

    return defaults


# ---------------------------------------------------------------------------
# Nutrientes e ingredientes
# ---------------------------------------------------------------------------

def get_nutrient_names(product_data: Dict[str, Any]) -> List[str]:
    df = product_data.get("nutrients", pd.DataFrame())
    if df is None or df.empty or "Nutriente" not in df.columns:
        return []
    return unique_clean(df["Nutriente"].tolist())


def default_nutrients(available: List[str]) -> List[str]:
    if not available:
        return []
    selected = []
    available_by_key = {norm_key(v): v for v in available}
    for hint in DEFAULT_NUTRIENT_HINTS:
        hint_key = norm_key(hint)
        for key, val in available_by_key.items():
            if hint_key in key or key in hint_key:
                if val not in selected:
                    selected.append(val)
                break
    return selected[:12] if selected else available[:10]


def selected_nutrients_df(product_data: Dict[str, Any], selected: List[str]) -> pd.DataFrame:
    df = product_data.get("nutrients", pd.DataFrame())
    if df is None or df.empty or "Nutriente" not in df.columns:
        return pd.DataFrame()
    selected_keys = {norm_key(x) for x in selected}
    out = df[df["Nutriente"].map(norm_key).isin(selected_keys)].copy()
    cols = [c for c in ["Nutriente", "Valor", "Restricción", "Mínimo", "Máximo"] if c in out.columns]
    return out[cols].reset_index(drop=True)


def limits_df(df: pd.DataFrame, kind: str) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    cols = []
    if kind == "nutrients":
        cols = [c for c in ["Nutriente", "Valor", "Restricción", "Mínimo", "Máximo"] if c in df.columns]
    else:
        cols = [c for c in ["Materia prima", "%", "Restricción", "Mínimo", "Máximo"] if c in df.columns]
    if not cols:
        return pd.DataFrame()
    return df[cols].copy().reset_index(drop=True)



def format_cell_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        if float(value).is_integer():
            return str(int(value))
        return f"{float(value):.4g}".replace(".", ",")
    return clean_text(value)


def dataframe_to_markdown_table(df: pd.DataFrame, columns: List[str]) -> str:
    if df is None or df.empty:
        return ""
    rows = []
    for _, row in df.iterrows():
        vals = [format_cell_value(row.get(col, "")) for col in columns]
        if any(vals):
            rows.append(vals)
    if not rows:
        return ""

    widths = [len(col) for col in columns]
    for vals in rows:
        for i, val in enumerate(vals):
            widths[i] = min(max(widths[i], len(val)), 42)

    def trim(val: str, width: int) -> str:
        if len(val) <= width:
            return val
        return val[: max(0, width - 1)] + "…"

    def fmt(vals: List[str]) -> str:
        return " | ".join(trim(vals[i], widths[i]).ljust(widths[i]) for i in range(len(columns)))

    lines = [fmt(columns), " | ".join("-" * w for w in widths)]
    lines.extend(fmt(vals) for vals in rows)
    return "\n".join(lines)


def table_to_text(df: pd.DataFrame, title: str = "") -> str:
    if df is None or df.empty:
        return ""

    cols = list(df.columns)
    norm_cols = {norm_key(c): c for c in cols}

    if "nutriente" in norm_cols:
        columns = [c for c in ["Nutriente", "Valor", "Mínimo", "Máximo", "Restricción"] if c in df.columns]
    elif "materiaprima" in norm_cols or "%" in cols:
        rename = {}
        if "Materia prima" in df.columns:
            rename["Materia prima"] = "Ingrediente"
        if "%" in df.columns:
            rename["%"] = "% fórmula"
        df = df.rename(columns=rename)
        columns = [c for c in ["Ingrediente", "% fórmula", "Mínimo", "Máximo", "Restricción"] if c in df.columns]
    else:
        columns = cols

    text = dataframe_to_markdown_table(df, columns)
    if title and text:
        return f"{title}\n{text}"
    return text

# ---------------------------------------------------------------------------
# Exportación
# ---------------------------------------------------------------------------


def is_table_separator_line(line: str) -> bool:
    parts = [part.strip() for part in line.split("|")]
    return len(parts) > 1 and all(part and set(part) <= {"-"} for part in parts)


def collect_markdown_table(lines: List[str], start: int) -> Tuple[Optional[List[List[str]]], int]:
    if start + 1 >= len(lines):
        return None, start
    header = lines[start]
    sep = lines[start + 1]
    if "|" not in header or not is_table_separator_line(sep):
        return None, start
    table_lines = [header, sep]
    i = start + 2
    while i < len(lines) and "|" in lines[i].strip() and lines[i].strip():
        table_lines.append(lines[i])
        i += 1
    rows = []
    for line in [table_lines[0]] + table_lines[2:]:
        rows.append([cell.strip() for cell in line.split("|")])
    return rows, i


def text_to_docx_bytes(text: str, title: str) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx no está instalado.")
    doc = Document()
    doc.add_heading(title, level=1)
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        cleaned = lines[i].strip()
        if not cleaned:
            doc.add_paragraph("")
            i += 1
            continue

        table_rows, next_i = collect_markdown_table(lines, i)
        if table_rows:
            table = doc.add_table(rows=1, cols=len(table_rows[0]))
            table.style = "Table Grid"
            for j, value in enumerate(table_rows[0]):
                table.rows[0].cells[j].text = value
            for row_values in table_rows[1:]:
                cells = table.add_row().cells
                for j, value in enumerate(row_values[:len(cells)]):
                    cells[j].text = value
            i = next_i
            continue

        if cleaned.isupper() and len(cleaned) < 90:
            doc.add_heading(cleaned.title(), level=2)
        elif cleaned.startswith("- "):
            doc.add_paragraph(cleaned[2:], style="List Bullet")
        else:
            doc.add_paragraph(cleaned)
        i += 1

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def text_to_pdf_bytes(text: str, title: str) -> bytes:
    if SimpleDocTemplate is None:
        raise RuntimeError("reportlab no está instalado.")
    bio = io.BytesIO()
    doc = SimpleDocTemplate(
        bio,
        pagesize=A4,
        rightMargin=1.4 * cm,
        leftMargin=1.4 * cm,
        topMargin=1.3 * cm,
        bottomMargin=1.3 * cm,
        title=title,
    )
    styles = getSampleStyleSheet()
    story = [Paragraph(xml_escape(title), styles["Title"]), Spacer(1, 0.35 * cm)]
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        cleaned = lines[i].strip()
        if not cleaned:
            story.append(Spacer(1, 0.16 * cm))
            i += 1
            continue

        table_rows, next_i = collect_markdown_table(lines, i)
        if table_rows and Table is not None:
            table_data = [[Paragraph(xml_escape(cell), styles["BodyText"]) for cell in row] for row in table_rows]
            table = Table(table_data, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EF2533")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(table)
            story.append(Spacer(1, 0.18 * cm))
            i = next_i
            continue

        if cleaned.isupper() and len(cleaned) < 90:
            story.append(Paragraph(xml_escape(cleaned.title()), styles["Heading2"]))
        else:
            story.append(Paragraph(xml_escape(cleaned), styles["BodyText"]))
            story.append(Spacer(1, 0.08 * cm))
        i += 1
    doc.build(story)
    return bio.getvalue()


def flat_to_xlsx_bytes(flat_rows: List[Dict[str, Any]], text: str, title: str) -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df = pd.DataFrame(flat_rows or [{"Campo": "Texto", "Valor": text}])
        df.to_excel(writer, sheet_name="Ficha", index=False)
        pd.DataFrame({"Texto final": text.splitlines()}).to_excel(writer, sheet_name="Texto", index=False)
        canva_df = build_canva_records(flat_rows)
        if not canva_df.empty:
            canva_df.to_excel(writer, sheet_name="Canva_Comercial", index=False)
        meta = pd.DataFrame(
            [
                {"Campo": "Título", "Valor": title},
                {"Campo": "Fecha exportación", "Valor": date.today().isoformat()},
                {"Campo": "Uso", "Valor": "La hoja Canva_Comercial está preparada para Canva Bulk Create / Google Sheets."},
            ]
        )
        meta.to_excel(writer, sheet_name="Metadatos", index=False)
    return bio.getvalue()


def batch_to_xlsx_bytes(batch_rows: List[Dict[str, Any]]) -> bytes:
    bio = io.BytesIO()
    df = pd.DataFrame(batch_rows)
    canva_df = build_canva_records(batch_rows)
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Acumulado", index=False)
        if not canva_df.empty:
            canva_df.to_excel(writer, sheet_name="Canva_Comercial", index=False)
        pd.DataFrame([
            {"Campo": "Fecha exportación", "Valor": date.today().isoformat()},
            {"Campo": "Nota", "Valor": "Canva_Comercial contiene una fila por producto/ficha con columnas simples para maquetación masiva."},
        ]).to_excel(writer, sheet_name="Metadatos", index=False)
    return bio.getvalue()

# ---------------------------------------------------------------------------
# Construcción de documentos
# ---------------------------------------------------------------------------


def build_context(
    product_display: str,
    product_data: Dict[str, Any],
    manual_values: Dict[str, str],
    selected_nutrients: Any,
) -> Dict[str, Any]:
    ctx = dict(manual_values)
    ctx.setdefault("Producto", product_display)
    ctx.setdefault("Nombre comercial", product_data.get("product_name", product_display))
    ctx.setdefault("Código / Specification", product_data.get("product_id", ""))
    ctx.setdefault("Coste por tonelada", product_data.get("cost_per_tonne", ""))
    ctx.setdefault("Fuente formulación", product_data.get("source_type", ""))

    if isinstance(selected_nutrients, dict):
        commercial_selected = selected_nutrients.get("commercial", [])
        technical_selected = selected_nutrients.get("technical", [])
    else:
        commercial_selected = selected_nutrients or []
        technical_selected = selected_nutrients or []

    commercial_df = selected_nutrients_df(product_data, commercial_selected)
    technical_df = selected_nutrients_df(product_data, technical_selected)
    ctx["Nutrientes seleccionados - FT Comercial"] = table_to_text(commercial_df)
    ctx["Nutrientes seleccionados - FT Técnica"] = table_to_text(technical_df)
    ctx["Nutrientes seleccionados"] = ctx["Nutrientes seleccionados - FT Técnica"] or ctx["Nutrientes seleccionados - FT Comercial"]

    nutrient_limits = limits_df(product_data.get("nutrient_limits", pd.DataFrame()), "nutrients")
    ingredient_limits = limits_df(product_data.get("ingredient_limits", pd.DataFrame()), "ingredients")
    ingredients = limits_df(product_data.get("ingredients", pd.DataFrame()), "ingredients")

    ctx["Límites de nutrientes"] = table_to_text(nutrient_limits)
    ctx["Límites de ingredientes"] = table_to_text(ingredient_limits)
    ctx["Ingredientes de fórmula"] = table_to_text(ingredients)

    return ctx

def field_value(ctx: Dict[str, Any], field: str) -> str:
    return clean_text(ctx.get(field, ""))



def build_document_text(
    doc_type: str,
    ctx: Dict[str, Any],
    selected_fields: List[str],
    include_nutrients: bool,
    include_limits: bool,
) -> Tuple[str, List[Dict[str, Any]]]:
    product = field_value(ctx, "Nombre comercial") or field_value(ctx, "Producto")
    title = f"{doc_type} - {product}".strip(" -")

    lines: List[str] = [title.upper(), ""]
    flat_rows: List[Dict[str, Any]] = [
        {"Tipo documento": doc_type, "Producto": product, "Campo": "Título", "Valor": title},
        {"Tipo documento": doc_type, "Producto": product, "Campo": "Producto", "Valor": field_value(ctx, "Producto")},
    ]

    for field in selected_fields:
        value = field_value(ctx, field)
        if not value:
            continue
        lines.append(field.upper())
        lines.append(value)
        lines.append("")
        flat_rows.append({"Tipo documento": doc_type, "Producto": product, "Campo": field, "Valor": value})

    if include_nutrients:
        profile = DOC_NUTRIENT_PROFILE.get(doc_type, "technical")
        nutrient_field = "Nutrientes seleccionados - FT Comercial" if profile == "commercial" else "Nutrientes seleccionados - FT Técnica"
        nutrient_value = field_value(ctx, nutrient_field) or field_value(ctx, "Nutrientes seleccionados")
        if nutrient_value:
            lines.append("NUTRIENTES")
            lines.append(nutrient_value)
            lines.append("")
            flat_rows.append({
                "Tipo documento": doc_type,
                "Producto": product,
                "Campo": "Nutrientes seleccionados",
                "Valor": nutrient_value,
            })

    if include_limits:
        if field_value(ctx, "Límites de nutrientes"):
            lines.append("LÍMITES DE NUTRIENTES")
            lines.append(field_value(ctx, "Límites de nutrientes"))
            lines.append("")
            flat_rows.append({
                "Tipo documento": doc_type,
                "Producto": product,
                "Campo": "Límites de nutrientes",
                "Valor": field_value(ctx, "Límites de nutrientes"),
            })
        if field_value(ctx, "Límites de ingredientes"):
            lines.append("LÍMITES DE INGREDIENTES")
            lines.append(field_value(ctx, "Límites de ingredientes"))
            lines.append("")
            flat_rows.append({
                "Tipo documento": doc_type,
                "Producto": product,
                "Campo": "Límites de ingredientes",
                "Valor": field_value(ctx, "Límites de ingredientes"),
            })

    return "\n".join(lines).strip() + "\n", flat_rows

def doc_defaults(doc_type: str) -> Tuple[List[str], bool, bool]:
    if doc_type == "Etiqueta":
        return LABEL_FIELDS, False, False
    if doc_type == "FT Calidad-Operaciones":
        return QUALITY_OPERATION_FIELDS, True, False
    if doc_type == "FT Comercial":
        return COMMERCIAL_FIELDS, True, False
    if doc_type == "FT Especificaciones":
        return SPECIFICATION_FIELDS, True, True
    return ALL_FIELD_OPTIONS, False, False


# ---------------------------------------------------------------------------
# UI Streamlit
# ---------------------------------------------------------------------------


def init_state() -> None:
    st.session_state.setdefault("selected_nutrients_global", [])
    st.session_state.setdefault("generated_docs", {})
    st.session_state.setdefault("batch_rows", [])
    st.session_state.setdefault("saved_nutrient_defaults", load_nutrient_defaults())
    st.session_state.setdefault("quality_operation_defaults", load_quality_operation_defaults())


def reset_current_product_state() -> None:
    preserve = {
        "selected_nutrients_global",
        "batch_rows",
        "saved_nutrient_defaults",
        "quality_operation_defaults",
        "select_nutrients_commercial",
        "select_nutrients_technical",
    }
    for key in list(st.session_state.keys()):
        if key in preserve:
            continue
        if key.startswith(("manual_", "select_", "generated_", "final_text_", "fields_")):
            del st.session_state[key]
    st.session_state["generated_docs"] = {}


def sidebar_uploads() -> Tuple[Dict[str, Dict[str, Any]], pd.DataFrame, pd.DataFrame]:
    st.sidebar.header("1. Carga de ficheros")
    template_status_sidebar()
    formula_file = st.sidebar.file_uploader(
        "Fichero de formulación",
        type=["xlsx", "xlsm", "xls", "xlsb", "ods", "csv", "tsv", "txt", "dat", "prn", "json", "pdf", "docx", "parquet"],
        help="Admite Excel/ODS, CSV/TSV/TXT, JSON, PDF, DOCX y Parquet. Para salidas Multi-Mix en texto detecta productos por 'Specification:' y por bloques 'SP:'.",
    )

    products: Dict[str, Dict[str, Any]] = {}
    etiquetas = pd.DataFrame()
    beneficios = pd.DataFrame()

    if formula_file is not None:
        with st.spinner("Leyendo formulación..."):
            products = load_formula_file(formula_file)
        st.sidebar.success(f"Productos detectados: {len(products)}" if products else "No se detectaron productos.")

    use_upload = False
    if MASTER_TEXTS_PATH.exists():
        st.sidebar.success(f"Textos maestros cargados desde {MASTER_TEXTS_PATH.as_posix()}")
        use_upload = st.sidebar.checkbox("Usar otro Excel de etiquetas y beneficios", value=False)
        if not use_upload:
            try:
                etiquetas, beneficios = load_parameter_file(MASTER_TEXTS_PATH)
                st.sidebar.caption(f"Etiquetas: {len(etiquetas)} · Beneficios: {len(beneficios)}")
            except Exception as exc:
                st.sidebar.error(f"No se pudo leer el Excel maestro: {exc}")
                use_upload = True
    else:
        st.sidebar.info("No se encuentra data/Etiquetas_y_beneficios_para_App_animales_destino.xlsx. Carga el Excel manualmente.")
        use_upload = True

    if use_upload:
        param_file = st.sidebar.file_uploader(
            "Excel de etiquetas y beneficios",
            type=["xlsx", "xlsm", "xls", "xlsb", "ods"],
            help="Debe contener hojas equivalentes a 'Etiquetas' y 'Beneficios'.",
        )
        if param_file is not None:
            with st.spinner("Leyendo textos parametrizados..."):
                etiquetas, beneficios = load_parameter_file(param_file)
            st.sidebar.success(
                f"Etiquetas: {len(etiquetas)} filas · Beneficios: {len(beneficios)} filas"
            )

    return products, etiquetas, beneficios

def select_base_fields(
    product_display: str,
    product_data: Dict[str, Any],
    etiquetas: pd.DataFrame,
    beneficios: pd.DataFrame,
) -> Dict[str, str]:
    st.header("2. Producto y textos base")

    defaults = infer_defaults_from_product(product_data.get("product_name", product_display), etiquetas, beneficios)

    species_options = unique_clean(get_unique(etiquetas, "Especie") + get_unique(beneficios, "Especie"))
    if not species_options:
        species_options = [""]

    species_default = defaults.get("Especie") if defaults.get("Especie") in species_options else species_options[0]
    species = st.selectbox(
        "Especie",
        species_options,
        index=species_options.index(species_default) if species_default in species_options else 0,
        key="select_especie",
    )

    subspecies_options = unique_clean(
        get_unique(filter_df(etiquetas, Especie=species), "Subespecie") +
        get_unique(filter_df(beneficios, Especie=species), "Subespecie")
    )
    if not subspecies_options:
        subspecies_options = [""]

    subspecies_default = defaults.get("Subespecie") if defaults.get("Subespecie") in subspecies_options else subspecies_options[0]
    subspecies = st.selectbox(
        "Subespecie",
        subspecies_options,
        index=subspecies_options.index(subspecies_default) if subspecies_default in subspecies_options else 0,
        key="select_subespecie",
    )

    lifestage_options = get_unique(
        filter_df(beneficios, Especie=species, Subespecie=subspecies),
        "Lifestage",
    )
    if not lifestage_options:
        lifestage_options = [""]

    lifestage_default = defaults.get("Lifestage") if defaults.get("Lifestage") in lifestage_options else lifestage_options[0]
    lifestage = st.selectbox(
        "Lifestage / Estado productivo",
        lifestage_options,
        index=lifestage_options.index(lifestage_default) if lifestage_default in lifestage_options else 0,
        key="select_lifestage",
    )

    # Etiqueta row selection.
    label_df = filter_df(etiquetas, Especie=species, Subespecie=subspecies)
    label_options = {}
    if not label_df.empty:
        for idx, row in label_df.iterrows():
            label_options[str(idx)] = option_label(
                row,
                ["Opción", "Animales de destino", "Tipo de pienso", "Modo de empleo", "Precauciones de uso"],
                max_len=180,
            )

    selected_label_idx = None
    if label_options:
        selected_label_idx = st.selectbox(
            "Opción de etiqueta propuesta",
            list(label_options.keys()),
            format_func=lambda k: label_options[k],
            key="select_label_option",
        )
        label_row = label_df.loc[int(selected_label_idx)]
    else:
        st.info("No hay opciones de etiqueta filtradas para la especie/subespecie seleccionada.")
        label_row = pd.Series(dtype=object)

    benefit_df = filter_df(beneficios, Especie=species, Subespecie=subspecies, Lifestage=lifestage)
    benefit_options = {}
    if not benefit_df.empty:
        for idx, row in benefit_df.iterrows():
            benefit_options[str(idx)] = option_label(row, ["Opción", "Foco-beneficio", "Texto beneficio"], max_len=180)

    selected_benefit_idx = None
    if benefit_options:
        selected_benefit_idx = st.selectbox(
            "Opción de beneficio propuesta",
            list(benefit_options.keys()),
            format_func=lambda k: benefit_options[k],
            key="select_benefit_option",
        )
        benefit_row = benefit_df.loc[int(selected_benefit_idx)]
    else:
        st.info("No hay beneficios filtrados para la combinación seleccionada.")
        benefit_row = pd.Series(dtype=object)

    tipo_pienso = clean_text(label_row.get("Tipo de pienso", "")) if not label_row.empty else ""
    animales = clean_text(label_row.get("Animales de destino", "")) if not label_row.empty else ""
    modo = clean_text(label_row.get("Modo de empleo", "")) if not label_row.empty else ""
    precauciones = clean_text(label_row.get("Precauciones de uso", "")) if not label_row.empty else ""
    foco = clean_text(benefit_row.get("Foco-beneficio", "")) if not benefit_row.empty else clean_text(label_row.get("Foco-beneficio", ""))
    beneficio = clean_text(benefit_row.get("Texto beneficio", "")) if not benefit_row.empty else ""

    row_data = product_data.get("row", {}) or {}
    product_name = clean_text(product_data.get("product_name", product_display))
    init_quality_widget_defaults(row_data, product_data)

    proposed_values = {
        "manual_nombre_comercial": product_name,
        "manual_tipo_pienso": tipo_pienso,
        "manual_animales_destino": animales,
        "manual_modo_empleo": modo,
        "manual_precauciones": precauciones,
        "manual_foco_beneficio": foco,
        "manual_beneficios": beneficio,
    }
    if st.button("Pasar opciones elegidas a campos editables", use_container_width=True):
        for key, value in proposed_values.items():
            st.session_state[key] = clean_text(value)
        st.rerun()

    st.subheader("Campos editables principales")
    col1, col2 = st.columns(2)
    with col1:
        nombre = st.text_input("Nombre comercial", value=product_name, key="manual_nombre_comercial")
        tipo_pienso = st.text_input("Tipo de pienso", value=tipo_pienso, key="manual_tipo_pienso")
        animales = st.text_area("Animales de destino", value=animales, height=90, key="manual_animales_destino")
    with col2:
        foco = st.text_area("Foco-beneficio", value=foco, height=90, key="manual_foco_beneficio")
        beneficio = st.text_area("Beneficios", value=beneficio, height=130, key="manual_beneficios")

    modo = st.text_area("Modo de empleo", value=modo, height=130, key="manual_modo_empleo")
    precauciones = st.text_area("Precauciones de uso / recomendaciones de manejo", value=precauciones, height=100, key="manual_precauciones")

    with st.expander("Campos adicionales de Calidad-Operaciones y Especificaciones", expanded=False):
        c1, c2 = st.columns(2)
        with c1:
            fecha = st.text_input("Fecha", value=date.today().isoformat(), key="manual_fecha")
            version = st.text_input("Versión", value=clean_text(row_data.get("Versión", row_data.get("Version", "1"))), key="manual_version")
            codigo_unite = st.text_input("Código Unite", value=clean_text(row_data.get("Código Unite", row_data.get("Codigo Unite", product_data.get("product_id", "")))), key="manual_codigo_unite")
            ean = st.text_input("EAN", value=clean_text(row_data.get("EAN", "")), key="manual_ean")
            presentacion = st.text_input("Presentación", value=clean_text(row_data.get("Presentación", row_data.get("PRESENTACION", ""))), key="manual_presentacion")
            peso_saco = st.text_input("Peso del producto en saco", value=clean_text(row_data.get("Peso del producto en saco", "")), key="manual_peso_saco")
            homologado = st.text_input("Homologado", value=clean_text(row_data.get("Homologado", row_data.get("HOMOLOGADO", ""))), key="manual_homologado")
            medicado = st.text_input("Medicado", value=clean_text(row_data.get("Medicado", row_data.get("MEDICADO", ""))), key="manual_medicado")
            periodo_espera = st.text_input("Periodo de espera", value=clean_text(row_data.get("Periodo de espera", "")), key="manual_periodo_espera")
            fabrica = st.text_input("Fábrica", value=clean_text(row_data.get("Fábrica", row_data.get("FABRICA", ""))), key="manual_fabrica")
            direccion = st.text_input("Dirección", value=clean_text(row_data.get("Dirección", row_data.get("DIRECCION", ""))), key="manual_direccion")
        with c2:
            codigo_postal = st.text_input("Código postal", value=clean_text(row_data.get("Código postal", row_data.get("CODIGO POSTAL", ""))), key="manual_codigo_postal")
            poblacion = st.text_input("Población", value=clean_text(row_data.get("Población", row_data.get("POBLACION", ""))), key="manual_poblacion")
            provincia = st.text_input("Provincia", value=clean_text(row_data.get("Provincia", row_data.get("PROVINCIA", ""))), key="manual_provincia")
            telefono = st.text_input("Teléfono", value=clean_text(row_data.get("Teléfono", row_data.get("TELEFONO", ""))), key="manual_telefono")
            imagen_envase = st.text_input("Imagen envase", value=clean_text(row_data.get("Imagen envase", "")), key="manual_imagen_envase")
            ficha_envase = st.text_input("Ficha técnica del envase", value=clean_text(row_data.get("Ficha técnica del envase", "")), key="manual_ficha_envase")
            formato_pale = st.text_area("Formato de palé. Mosaico, alturas y peso", value=clean_text(row_data.get("Formato de palé. Mosaico y alturas y peso", "")), height=70, key="manual_formato_pale")
            dur_min = st.text_input("Durabilidad mínima %", value=clean_text(row_data.get("Durabilidad minima %", row_data.get("Durabilidad mínima %", ""))), key="manual_dur_min")
            dur_alert = st.text_input("Alerta durabilidad %", value=clean_text(row_data.get("Alerta durabilidad %", "")), key="manual_dur_alert")
            finos_max = st.text_input("Finos máximos %", value=clean_text(row_data.get("Finos máximos %", "")), key="manual_finos_max")
            finos_alert = st.text_input("Alerta finos %", value=clean_text(row_data.get("Alerta finos %", "")), key="manual_finos_alert")

        definicion = st.text_area("Definición / Posicionamiento", value=clean_text(row_data.get("Definición / Posicionamiento", "")), height=100, key="manual_definicion")
        caracteristicas = st.text_area("Características", value=clean_text(row_data.get("Características", "")), height=100, key="manual_caracteristicas")
        caracteristicas_nut = st.text_area("Características nutricionales", value=clean_text(row_data.get("Características nutricionales", "")), height=90, key="manual_caracteristicas_nut")
        materias = st.text_area("Materias primas, aditivos y correctores específicos", value=clean_text(row_data.get("Materias primas, aditivos, correctores ESPECIFICOS", "")), height=90, key="manual_materias")
        analiticas = st.text_area("Analíticas especiales además del plan analítico", value=clean_text(row_data.get("Analíticas especiales además de las ya definidas por plan analítico.", "")), height=70, key="manual_analiticas")

        if st.button("Guardar campos adicionales por defecto", key="save_quality_operation_defaults"):
            ok, msg = save_quality_operation_defaults(collect_quality_values_from_session())
            if ok:
                st.success(msg)
            else:
                st.error(msg)

    values = {
        "Producto": product_display,
        "Nombre comercial": nombre,
        "Fecha": fecha,
        "Versión": version,
        "Código Unite": codigo_unite,
        "EAN": ean,
        "Tipo de pienso": tipo_pienso,
        "Especie": species,
        "Subespecie": subspecies,
        "Lifestage": lifestage,
        "Animales de destino": animales,
        "Modo de empleo": modo,
        "Precauciones de uso": precauciones,
        "Recomendaciones de manejo en etiquetado": precauciones,
        "Foco-beneficio": foco,
        "Beneficios": beneficio,
        "Definición / Posicionamiento": definicion if "definicion" in locals() else "",
        "Características": caracteristicas if "caracteristicas" in locals() else "",
        "Imagen envase": imagen_envase if "imagen_envase" in locals() else "",
        "Peso del producto en saco": peso_saco if "peso_saco" in locals() else "",
        "Ficha técnica del envase": ficha_envase if "ficha_envase" in locals() else "",
        "Formato de palé. Mosaico, alturas y peso": formato_pale if "formato_pale" in locals() else "",
        "Características nutricionales": caracteristicas_nut if "caracteristicas_nut" in locals() else "",
        "Materias primas, aditivos y correctores específicos": materias if "materias" in locals() else "",
        "Analíticas especiales además del plan analítico": analiticas if "analiticas" in locals() else "",
        "Presentación": presentacion if "presentacion" in locals() else "",
        "Durabilidad mínima %": dur_min if "dur_min" in locals() else "",
        "Alerta durabilidad %": dur_alert if "dur_alert" in locals() else "",
        "Finos máximos %": finos_max if "finos_max" in locals() else "",
        "Alerta finos %": finos_alert if "finos_alert" in locals() else "",
        "Homologado": homologado if "homologado" in locals() else "",
        "Medicado": medicado if "medicado" in locals() else "",
        "Periodo de espera": periodo_espera if "periodo_espera" in locals() else "",
        "Fábrica": fabrica if "fabrica" in locals() else "",
        "Dirección": direccion if "direccion" in locals() else "",
        "Código postal": codigo_postal if "codigo_postal" in locals() else "",
        "Población": poblacion if "poblacion" in locals() else "",
        "Provincia": provincia if "provincia" in locals() else "",
        "Teléfono": telefono if "telefono" in locals() else "",
    }
    return values



def nutrient_selectors(product_data: Dict[str, Any]) -> Dict[str, List[str]]:
    st.header("3. Nutrientes")
    available = get_nutrient_names(product_data)

    if not available:
        st.info("No se han detectado nutrientes numéricos en el producto seleccionado.")
        return {"commercial": [], "technical": []}

    selections: Dict[str, List[str]] = {}
    c1, c2 = st.columns(2)
    selector_specs = [
        ("commercial", "FT Comercial", "select_nutrients_commercial", c1),
        ("technical", "FT Calidad-Operaciones y FT Especificaciones", "select_nutrients_technical", c2),
    ]
    for profile, label, key, col in selector_specs:
        with col:
            default = default_nutrients_for_profile(profile, available, key)
            selected = st.multiselect(
                f"Nutrientes para {label}",
                options=available,
                default=default,
                key=key,
                help="Puede ser distinto por tipo de ficha. Use Guardar nutrientes para reutilizar esta selección por defecto.",
            )
            selections[profile] = selected
            if st.button(f"Guardar nutrientes - {label}", key=f"save_nutrients_{profile}"):
                ok, msg = set_saved_nutrient_defaults(profile, selected)
                if ok:
                    st.success(msg)
                else:
                    st.error(msg)

    with st.expander("Vista de nutrientes seleccionados", expanded=False):
        if selections.get("commercial"):
            st.write("**FT Comercial**")
            st.text(table_to_text(selected_nutrients_df(product_data, selections["commercial"])))
        if selections.get("technical"):
            st.write("**FT Calidad-Operaciones y FT Especificaciones**")
            st.text(table_to_text(selected_nutrients_df(product_data, selections["technical"])))

    return selections

def field_selectors() -> Dict[str, Dict[str, Any]]:
    st.header("4. Generación de documentos")
    selections: Dict[str, Dict[str, Any]] = {}
    tabs = st.tabs(DOC_TYPES)
    for doc_type, tab in zip(DOC_TYPES, tabs):
        with tab:
            defaults, include_nutrients, include_limits = doc_defaults(doc_type)
            selected_fields = st.multiselect(
                f"Campos a incluir en {doc_type}",
                options=ALL_FIELD_OPTIONS,
                default=[x for x in defaults if x in ALL_FIELD_OPTIONS],
                key=f"fields_{doc_type}",
            )
            selections[doc_type] = {
                "fields": selected_fields,
                "include_nutrients": include_nutrients,
                "include_limits": include_limits,
            }
    return selections



def render_downloads(
    doc_type: str,
    title_base: str,
    generated_text: str,
    flat_rows: List[Dict[str, Any]],
) -> None:
    text_key = f"final_text_{doc_type}"
    edited_text = st.text_area(
        f"Texto final editable - {doc_type}",
        value=st.session_state.get(text_key, generated_text),
        height=360,
        key=text_key,
    )

    file_base = safe_filename(f"{doc_type}_{title_base}")
    fields = flat_rows_to_field_map(flat_rows)
    template_options = template_options_for_doc(doc_type, fields)

    selected_template = st.selectbox(
        "Plantilla HTML/PDF",
        options=template_options,
        format_func=lambda name: TEMPLATE_DISPLAY_NAMES.get(name, name),
        key=f"html_template_{doc_type}",
        help="Automática usa la plantilla recomendada según tipo de ficha y especie. También puede forzarse manualmente.",
    )

    html, template_name, cover_file = render_nanta_html(
        doc_type,
        f"{doc_type} - {title_base}",
        flat_rows,
        edited_text,
        selected_template=selected_template,
    )
    st.caption(f"Plantilla aplicada: {TEMPLATE_DISPLAY_NAMES.get(template_name, template_name)} · Imagen: {cover_file}")

    with st.expander("Vista previa maquetada", expanded=True):
        components.html(html, height=900, scrolling=True)

    c1, c2, c3, c4, c5, c6 = st.columns(6)
    with c1:
        st.download_button(
            "TXT",
            edited_text.encode("utf-8"),
            file_name=f"{file_base}.txt",
            mime="text/plain",
            key=f"download_txt_{doc_type}",
        )
    with c2:
        try:
            docx_bytes = text_to_docx_bytes(edited_text, f"{doc_type} - {title_base}")
            st.download_button(
                "DOCX",
                docx_bytes,
                file_name=f"{file_base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_docx_{doc_type}",
            )
        except Exception as exc:
            st.caption(f"DOCX no disponible: {exc}")
    with c3:
        try:
            pdf_bytes = text_to_pdf_bytes(edited_text, f"{doc_type} - {title_base}")
            st.download_button(
                "PDF simple",
                pdf_bytes,
                file_name=f"{file_base}.pdf",
                mime="application/pdf",
                key=f"download_pdf_{doc_type}",
                help="PDF técnico simple. Para la ficha visual, descarga HTML/PDF NANTA y guárdalo como PDF desde el navegador.",
            )
        except Exception as exc:
            st.caption(f"PDF no disponible: {exc}")
    with c4:
        st.download_button(
            "HTML/PDF NANTA",
            html.encode("utf-8"),
            file_name=f"{file_base}_NANTA.html",
            mime="text/html",
            key=f"download_html_{doc_type}",
            help="Ficha maquetada con la plantilla elegida. Abrir en navegador e imprimir/guardar como PDF.",
        )
    with c5:
        xlsx_bytes = flat_to_xlsx_bytes(flat_rows, edited_text, f"{doc_type} - {title_base}")
        st.download_button(
            "Excel",
            xlsx_bytes,
            file_name=f"{file_base}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"download_xlsx_{doc_type}",
            help="Incluye hoja Canva_Comercial.",
        )
    with c6:
        if st.button("Acumular", key=f"batch_{doc_type}"):
            rows = [dict(row, Texto_final=edited_text, Plantilla_HTML=template_name) for row in flat_rows]
            st.session_state["batch_rows"].extend(rows)
            st.success(f"{doc_type} añadido al acumulado.")

def accumulated_export_panel() -> None:
    st.sidebar.header("Acumulado")
    n = len(st.session_state.get("batch_rows", []))
    st.sidebar.write(f"Filas acumuladas: {n}")
    if n:
        xlsx_bytes = batch_to_xlsx_bytes(st.session_state["batch_rows"])
        st.sidebar.download_button(
            "Descargar Excel acumulado",
            xlsx_bytes,
            file_name="acumulado_fichas_piensos.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="download_batch",
        )
        if st.sidebar.button("Vaciar acumulado"):
            st.session_state["batch_rows"] = []
            st.rerun()


def main() -> None:
    st.set_page_config(page_title=APP_TITLE, page_icon="🧾", layout="wide")
    init_state()

    st.title(APP_TITLE)
    st.caption("Prototipo operativo para generación rápida de etiquetas y fichas técnicas a partir de formulación y textos parametrizados.")

    products, etiquetas, beneficios = sidebar_uploads()
    accumulated_export_panel()

    if not products:
        st.info("Carga primero un fichero de formulación. Después se habilitará el selector de producto.")
        with st.expander("Formatos de formulación admitidos"):
            st.write(
                """
                - Excel/ODS con tablas de productos.
                - Excel de una columna con salida textual tipo Multi-Mix.
                - CSV/TSV/TXT/DAT/PRN.
                - JSON.
                - PDF y DOCX con texto extraíble.
                - Parquet.
                """
            )
        return

    st.sidebar.header("2. Producto")
    product_display = st.sidebar.selectbox("Producto", options=list(products.keys()), key="select_product")
    product_data = products[product_display]

    if st.sidebar.button("Borrar información del producto actual"):
        reset_current_product_state()
        st.rerun()

    with st.expander("Vista rápida del producto detectado", expanded=False):
        st.write(f"**Nombre:** {product_data.get('product_name', '')}")
        st.write(f"**Código/Specification:** {product_data.get('product_id', '')}")
        st.write(f"**Tipo de origen:** {product_data.get('source_type', '')}")
        nutrients = product_data.get("nutrients", pd.DataFrame())
        ingredients = product_data.get("ingredients", pd.DataFrame())
        if nutrients is not None and not nutrients.empty:
            st.write("Nutrientes detectados")
            st.dataframe(nutrients.head(50), use_container_width=True)
        if ingredients is not None and not ingredients.empty:
            st.write("Ingredientes detectados")
            st.dataframe(ingredients.head(50), use_container_width=True)

    manual_values = select_base_fields(product_display, product_data, etiquetas, beneficios)
    selected_nutrients = nutrient_selectors(product_data)
    doc_selections = field_selectors()

    ctx = build_context(product_display, product_data, manual_values, selected_nutrients)
    product_title = field_value(ctx, "Nombre comercial") or product_display

    c1, c2, c3, c4 = st.columns(4)
    button_map = {
        "Etiqueta": c1,
        "FT Calidad-Operaciones": c2,
        "FT Comercial": c3,
        "FT Especificaciones": c4,
    }

    for doc_type, col in button_map.items():
        with col:
            if st.button(f"Generar {doc_type}", key=f"generate_{doc_type}", use_container_width=True):
                sel = doc_selections[doc_type]
                text, flat = build_document_text(
                    doc_type,
                    ctx,
                    sel["fields"],
                    include_nutrients=sel["include_nutrients"],
                    include_limits=sel["include_limits"],
                )
                st.session_state["generated_docs"][doc_type] = {"text": text, "flat": flat}

    if st.session_state["generated_docs"]:
        st.subheader("5. Resultado y descargas")
        result_tabs = st.tabs(list(st.session_state["generated_docs"].keys()))
        for doc_type, tab in zip(st.session_state["generated_docs"].keys(), result_tabs):
            with tab:
                data = st.session_state["generated_docs"][doc_type]
                render_downloads(doc_type, product_title, data["text"], data["flat"])
    else:
        st.info("Selecciona campos y pulsa uno de los cuatro botones de generación.")


if __name__ == "__main__":
    main()
